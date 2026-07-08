"""Phase 4 — research report generation.

Produces:
  draft_interim_report(study_id)   → .docx (Background / Methods /
                                     Patient flow / Table 1 baseline /
                                     Safety / Efficacy KM / Discussion)
  draft_table_1(study_id)          → Table 1 baseline characteristics
  generate_consort_diagram(study_id) → CONSORT-style SVG
  generate_km_curve(study_id, endpoint) → KM curve PNG
  export_cohort_xlsx(study_id, deidentify=True) → .xlsx for analysis
"""
from __future__ import annotations

import json
import logging
import os
import sqlite3
import tempfile
import time
import uuid
from typing import Optional

from nexus_server.database import get_db_connection
from nexus_server.research.patient_facts import get_patient_facts

logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────


def _study(conn: sqlite3.Connection, user_id: str, study_id: str) -> dict:
    r = conn.execute(
        "SELECT display_name, short_code, phase, target_n, primary_endpoint, "
        "secondary_endpoints_json, protocol_summary "
        "FROM research_studies WHERE user_id = ? AND study_id = ?",
        (user_id, study_id),
    ).fetchone()
    if not r:
        raise RuntimeError(f"unknown study {study_id}")
    return dict(
        display_name=r[0], short_code=r[1], phase=r[2], target_n=r[3],
        primary_endpoint=r[4],
        secondary_endpoints=json.loads(r[5] or "[]"),
        protocol_summary=r[6],
    )


def _enrolled_patients(conn: sqlite3.Connection,
                       user_id: str, study_id: str) -> list[tuple[str, int, str]]:
    rows = conn.execute(
        "SELECT patient_hash, enrollment_seq, arm FROM study_enrollments "
        "WHERE user_id = ? AND study_id = ? "
        "ORDER BY enrollment_seq ASC",
        (user_id, study_id),
    ).fetchall()
    return [(r[0], r[1], r[2]) for r in rows]


def _study_flow(conn: sqlite3.Connection,
                user_id: str, study_id: str) -> dict:
    """For CONSORT: counts at each stage."""
    rows = conn.execute(
        "SELECT decision, COUNT(*) FROM screening_evaluations "
        "WHERE user_id = ? AND study_id = ? GROUP BY decision",
        (user_id, study_id),
    ).fetchall()
    decision_counts = {r[0]: r[1] for r in rows}

    enrollment_rows = conn.execute(
        "SELECT status, COUNT(*) FROM study_enrollments "
        "WHERE user_id = ? AND study_id = ? GROUP BY status",
        (user_id, study_id),
    ).fetchall()
    enroll_counts = {r[0]: r[1] for r in enrollment_rows}

    return {
        "screened":   sum(decision_counts.values()),
        "excluded":   decision_counts.get("excluded", 0),
        "invited":    decision_counts.get("invited", 0),
        "enrolled":   enroll_counts.get("enrolled", 0),
        "withdrawn":  enroll_counts.get("withdrawn", 0),
        "completed":  enroll_counts.get("completed", 0),
    }


def _table_1(conn: sqlite3.Connection,
             user_id: str, study_id: str,
             deidentify: bool = True) -> list[dict]:
    """Per-patient baseline row."""
    out: list[dict] = []
    for ph, seq, arm in _enrolled_patients(conn, user_id, study_id):
        f = get_patient_facts(conn, user_id, ph)
        row = {
            "#": seq,
            "Arm": arm or "",
            "Age": f.age or "",
            "Sex": f.sex or "",
            "ECOG": f.ecog if f.ecog is not None else "",
            "Pathology": f.pathology or "",
            "Stage": f.ajcc_stage or f.valg_stage or "",
            "Driver": f.driver_mutation or "",
            "Prior lines": ", ".join(f.prior_lines[:3]),
        }
        if not deidentify:
            row["patient_hash"] = ph
        out.append(row)
    return out


# ─────────────────────────────────────────────────────────────────────
# Public API
# ─────────────────────────────────────────────────────────────────────


def draft_table_1(user_id: str, study_id: str,
                  deidentify: bool = True) -> list[dict]:
    with get_db_connection() as conn:
        return _table_1(conn, user_id, study_id, deidentify=deidentify)


def generate_consort_diagram(user_id: str, study_id: str) -> str:
    """Render a CONSORT-style SVG to a temp file. Returns the path."""
    with get_db_connection() as conn:
        flow = _study_flow(conn, user_id, study_id)
        study = _study(conn, user_id, study_id)
    svg = _render_consort_svg(study["display_name"], flow)
    path = os.path.join(tempfile.gettempdir(), f"consort-{study_id}-{uuid.uuid4().hex[:6]}.svg")
    with open(path, "w") as fp:
        fp.write(svg)
    return path


def _render_consort_svg(title: str, flow: dict) -> str:
    boxes = [
        ("Screened",   flow["screened"]),
        ("Excluded",   flow["excluded"]),
        ("Invited",    flow["invited"]),
        ("Enrolled",   flow["enrolled"]),
        ("Withdrawn",  flow["withdrawn"]),
        ("Completed",  flow["completed"]),
    ]
    h = 80
    height = h * len(boxes) + 100
    rects = []
    for i, (label, n) in enumerate(boxes):
        y = 60 + i * h
        rects.append(
            f'<rect x="80" y="{y}" width="280" height="50" '
            f'fill="#eef" stroke="#446" rx="6"/>'
            f'<text x="220" y="{y+30}" font-size="16" '
            f'text-anchor="middle" font-family="sans-serif">{label}: {n}</text>'
        )
        if i + 1 < len(boxes):
            rects.append(
                f'<line x1="220" y1="{y+50}" x2="220" y2="{y+60+h-50}" '
                f'stroke="#446" marker-end="url(#arr)"/>'
            )
    return (
        f'<svg xmlns="http://www.w3.org/2000/svg" width="440" height="{height}">'
        f'<defs><marker id="arr" markerWidth="10" markerHeight="10" '
        f'refX="5" refY="3" orient="auto"><path d="M0,0 L0,6 L5,3 z" '
        f'fill="#446"/></marker></defs>'
        f'<text x="220" y="30" font-size="20" font-weight="bold" '
        f'text-anchor="middle" font-family="sans-serif">{title}</text>'
        + "".join(rects) +
        "</svg>"
    )


def generate_km_curve(user_id: str, study_id: str,
                      endpoint: str = "PFS") -> Optional[str]:
    """Render a Kaplan-Meier curve PNG. Falls back to None if no data
    or matplotlib unavailable."""
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except ImportError:
        return None

    with get_db_connection() as conn:
        # Until we have explicit event data, derive PFS placeholder
        # from enrolled_at + best_response: PR/CR → not progressed,
        # PD → "progressed at last visit". This is a placeholder; real
        # PFS requires actual progression dates. For now we render
        # constant survival.
        rows = conn.execute(
            "SELECT enrolled_at, status FROM study_enrollments "
            "WHERE user_id=? AND study_id=?", (user_id, study_id),
        ).fetchall()

    if not rows:
        return None

    now_ms = int(time.time() * 1000)
    months_seen = sorted({max(0, (now_ms - r[0]) // (30 * 86400 * 1000)) for r in rows})
    n = len(rows)

    fig, ax = plt.subplots(figsize=(6, 4))
    # Placeholder: until we have progression dates, draw all-at-risk
    # constant survival = 100%. The y axis encoding is what matters.
    xs = list(range(0, max(months_seen + [12]) + 1))
    ys = [100.0] * len(xs)
    ax.step(xs, ys, where="post", label=endpoint)
    ax.set_xlabel("Months from enrollment")
    ax.set_ylabel("Survival (%)")
    ax.set_title(f"KM — {endpoint} (n={n}, placeholder)")
    ax.set_ylim(0, 105)
    ax.grid(True, alpha=0.3)
    ax.legend()
    path = os.path.join(tempfile.gettempdir(),
                        f"km-{study_id}-{endpoint}-{uuid.uuid4().hex[:6]}.png")
    fig.tight_layout()
    fig.savefig(path, dpi=120)
    plt.close(fig)
    return path


def export_cohort_xlsx(user_id: str, study_id: str,
                       deidentify: bool = True) -> str:
    """Write a deidentified .xlsx of baseline + events. Returns path."""
    try:
        from openpyxl import Workbook
    except ImportError:
        # Fallback: write CSV alongside .xlsx extension so the path
        # parameter still resolves.
        return _export_csv(user_id, study_id, deidentify)

    with get_db_connection() as conn:
        rows = _table_1(conn, user_id, study_id, deidentify=deidentify)
        flow = _study_flow(conn, user_id, study_id)

    wb = Workbook()
    ws = wb.active
    ws.title = "Baseline"
    if rows:
        headers = list(rows[0].keys())
        ws.append(headers)
        for r in rows:
            ws.append([r.get(h, "") for h in headers])

    flow_ws = wb.create_sheet("Flow")
    flow_ws.append(["Stage", "Count"])
    for k, v in flow.items():
        flow_ws.append([k, v])

    path = os.path.join(tempfile.gettempdir(),
                        f"cohort-{study_id}-{uuid.uuid4().hex[:6]}.xlsx")
    wb.save(path)
    return path


def _export_csv(user_id: str, study_id: str, deidentify: bool) -> str:
    import csv
    with get_db_connection() as conn:
        rows = _table_1(conn, user_id, study_id, deidentify=deidentify)
    path = os.path.join(tempfile.gettempdir(),
                        f"cohort-{study_id}-{uuid.uuid4().hex[:6]}.csv")
    if rows:
        with open(path, "w", newline="") as fp:
            w = csv.DictWriter(fp, fieldnames=list(rows[0].keys()))
            w.writeheader()
            for r in rows:
                w.writerow(r)
    else:
        open(path, "w").close()
    return path


def draft_interim_report(user_id: str, study_id: str) -> str:
    """Generate the structured interim report .docx. Returns file_id.

    Sections (per design §3.3.5 → §5.5):
      1. Background & objectives
      2. Methods
      3. Results — patient flow / Table 1 / Safety / Efficacy
      4. Discussion (skeleton)
    """
    try:
        from docx import Document
    except ImportError:
        # Fall back to markdown
        return _draft_interim_report_md(user_id, study_id)

    with get_db_connection() as conn:
        study = _study(conn, user_id, study_id)
        rows  = _table_1(conn, user_id, study_id, deidentify=True)
        flow  = _study_flow(conn, user_id, study_id)

    doc = Document()
    doc.add_heading(f"Interim report — {study['display_name']}", level=0)
    doc.add_paragraph(f"Short code: {study['short_code']} · Phase {study['phase']}")
    doc.add_paragraph(f"Primary endpoint: {study['primary_endpoint'] or '—'}")
    doc.add_paragraph(f"Target N: {study['target_n']} · Enrolled: {flow['enrolled']}")
    doc.add_paragraph("")

    doc.add_heading("1. Background & objectives", level=1)
    doc.add_paragraph(study["protocol_summary"] or
                      "Protocol summary not provided.")

    doc.add_heading("2. Methods", level=1)
    doc.add_paragraph("Patients were screened via the Research Workspace "
                      "Eligibility engine (auto-rule + auto-llm + medic review). "
                      "Enrollment is recorded as the medic-confirmed event.")

    doc.add_heading("3. Results", level=1)
    doc.add_heading("3.1 Patient flow", level=2)
    for k in ("screened", "excluded", "invited", "enrolled",
              "withdrawn", "completed"):
        doc.add_paragraph(f"{k}: {flow.get(k, 0)}")
    doc.add_heading("3.2 Baseline (Table 1)", level=2)
    if rows:
        headers = list(rows[0].keys())
        tbl = doc.add_table(rows=1, cols=len(headers))
        for i, h in enumerate(headers):
            tbl.rows[0].cells[i].text = h
        for r in rows:
            cells = tbl.add_row().cells
            for i, h in enumerate(headers):
                cells[i].text = str(r.get(h, ""))
    else:
        doc.add_paragraph("No enrolled patients yet.")

    doc.add_heading("3.3 Safety", level=2)
    doc.add_paragraph("(Safety table — populated from study_observations "
                      "with ae_grade_confirmed=1 in Phase 4.5)")

    doc.add_heading("3.4 Efficacy", level=2)
    doc.add_paragraph("(KM curves placeholder — see /reports/km endpoint)")

    doc.add_heading("4. Discussion", level=1)
    doc.add_paragraph("To be drafted.")

    # Persist to uploads/ so the file is reachable from the
    # /uploads/{id} endpoint the desktop already uses.
    file_id = _persist_report_file(user_id, doc, prefix=f"interim-{study['short_code']}")
    return file_id


def _draft_interim_report_md(user_id: str, study_id: str) -> str:
    """Markdown fallback when python-docx isn't available."""
    with get_db_connection() as conn:
        study = _study(conn, user_id, study_id)
        rows  = _table_1(conn, user_id, study_id, deidentify=True)
        flow  = _study_flow(conn, user_id, study_id)
    md = [
        f"# Interim report — {study['display_name']}",
        f"Short code: {study['short_code']} · Phase {study['phase']}",
        f"Target N: {study['target_n']} · Enrolled: {flow['enrolled']}",
        "## Flow",
        *[f"- {k}: {v}" for k, v in flow.items()],
        "## Baseline (Table 1)",
        " | ".join(rows[0].keys()) if rows else "(no patients)",
    ]
    if rows:
        md.append(" | ".join("---" for _ in rows[0].keys()))
        for r in rows:
            md.append(" | ".join(str(v) for v in r.values()))
    return _persist_report_bytes(
        user_id, "\n".join(md).encode("utf-8"),
        prefix=f"interim-{study['short_code']}", ext=".md",
    )


def _persist_report_file(user_id: str, doc, *, prefix: str) -> str:
    """Save a python-docx Document to disk + register as an upload."""
    fid = f"report-{uuid.uuid4().hex[:12]}"
    path = os.path.join(tempfile.gettempdir(), f"{prefix}-{fid}.docx")
    doc.save(path)
    _register_upload(user_id, fid, path, mime=(
        "application/vnd.openxmlformats-officedocument."
        "wordprocessingml.document"))
    return fid


def _persist_report_bytes(
    user_id: str, data: bytes, *, prefix: str, ext: str,
) -> str:
    fid = f"report-{uuid.uuid4().hex[:12]}"
    path = os.path.join(tempfile.gettempdir(), f"{prefix}-{fid}{ext}")
    with open(path, "wb") as fp:
        fp.write(data)
    _register_upload(user_id, fid, path,
                     mime="text/markdown" if ext == ".md" else "application/octet-stream")
    return fid


def _register_upload(user_id: str, file_id: str, disk_path: str, *,
                     mime: str) -> None:
    """Best-effort: insert a row into uploads so the file is fetchable
    via /api/v1/files/{file_id} / /uploads/{file_id}."""
    from datetime import datetime, timezone
    name = os.path.basename(disk_path)
    size = os.path.getsize(disk_path)
    try:
        with get_db_connection() as conn:
            conn.execute(
                """
                INSERT INTO uploads
                (file_id, user_id, name, mime, size_bytes, disk_path,
                 created_at, sha256, gnfd_path)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (file_id, user_id, name, mime, size, disk_path,
                 datetime.now(timezone.utc).isoformat(), "", ""),
            )
            conn.commit()
    except sqlite3.Error as exc:
        logger.warning("could not register report upload: %s", exc)
