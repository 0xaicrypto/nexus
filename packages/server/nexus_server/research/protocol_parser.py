"""Protocol .docx → draft inclusion/exclusion/schedule (D7 batch-confirm).

Reads the uploaded .docx text + tables, runs heuristic + LLM-assisted
extraction, returns a draft for the batch-confirm UI. The caller
(``research_router.import_protocol``) does NOT auto-write the rules —
medic must hit Confirm in the UI (PATCH /studies/{id}).
"""
from __future__ import annotations

import json
import logging
import re
from typing import Optional

logger = logging.getLogger(__name__)


# Heuristic markers for Chinese clinical-trial protocols (the three
# reference protocols we've extracted from). English fallback included
# for ESMO/NCCN-style sub-protocols.
_INCL_HEADER = re.compile(
    r"(入[选组]标准|入[选组]条件|纳入标准|Inclusion\s+criteria)", re.I)
# "剔除" is just as common as "排除" in Chinese phase I/II protocols
# (剔除 = "strike out", "排除" = "exclude" — both refer to exclusion
# criteria in trial speak). Missing "剔除" here meant the parser stayed
# in inclusion-mode through the entire exclusion section, mis-tagging
# every exclusion item as inclusion. "不予入组" is the third common
# wording — typically used in IIT investigator-initiated trials.
_EXCL_HEADER = re.compile(
    r"(排除标准|排除条件|剔除标准|剔除条件|不予入组|Exclusion\s+criteria)", re.I)
_SCHEDULE_HEADER = re.compile(
    r"(随访|访视|时点|随访安排|Schedule|Follow-?up)", re.I)

# Any of these markers means the protocol moved on to a non-criterion
# section (treatment regimen, endpoints, statistics, CTCAE grading
# tables, etc.). Whatever section was active before resets to None so
# subsequent paragraphs don't get harvested into inclusion/exclusion.
# Without this guard the parser kept slurping protocol body text into
# whichever criteria section was last opened.
_END_OF_CRITERIA_HEADER = re.compile(
    r"^(治疗方案|研究方案|给药方案|疗效评定|疗效评价|终点指标|终点|"
    r"主要终点|次要终点|统计假设|统计方法|样本量|安全性监察|"
    r"安全性|不良事件|AE\s*分级|CTCAE|检查及对症治疗|会诊|数据管理|"
    r"伦理|知情同意书|研究流程|附录|参考文献|Treatment|Endpoints?|"
    r"Statistics?|Safety|Adverse\s+events?|References?)",
    re.I,
)


async def parse_protocol_docx(user_id: str, file_id: str) -> dict:
    """Extract a draft of inclusion + exclusion + schedule from a
    user-uploaded .docx whose file_id sits in the uploads table.

    Strategy (since 2026-06):
      1. Linearise the .docx (paragraphs + tables) into one markdown
         blob.
      2. Hand the WHOLE blob to an LLM with a structured-JSON contract
         and let it identify criteria + schedule + summary. The LLM
         knows that "剔除标准" means exclusion, that "疗效评定" is NOT
         a criterion, etc. — regex never will, no matter how many
         exceptions we accumulate.
      3. Fall back to the legacy regex parser if the LLM call fails
         (gateway missing, network down, schema violation). Server
         deployments without an LLM key still get *something*.

    Returns a dict the UI can show in the batch-confirm page:
      { "inclusion": [...], "exclusion": [...], "schedule": [...],
        "protocol_summary": "...", "notes": [warnings] }
    """
    # 1. Find the on-disk path of the uploaded file
    from nexus_server.database import get_db_connection
    with get_db_connection() as conn:
        row = conn.execute(
            "SELECT disk_path FROM uploads "
            "WHERE user_id = ? AND file_id = ?",
            (user_id, file_id),
        ).fetchone()
    if not row:
        return {"inclusion": [], "exclusion": [], "schedule": [],
                "protocol_summary": "",
                "notes": [f"upload {file_id} not found"]}
    disk_path = row[0]

    # 2. Extract raw text
    try:
        from docx import Document  # python-docx
    except ImportError:
        return {"inclusion": [], "exclusion": [], "schedule": [],
                "protocol_summary": "",
                "notes": ["python-docx not installed"]}

    # python-docx wraps both `FileNotFoundError` and `BadZipFile` into
    # the same opaque "Package not found at '<path>'" message, which
    # makes the resulting notes useless for debugging. Disambiguate up
    # front: existence + NFC-normalised lookup + magic-byte sniff.
    import unicodedata as _ud
    from pathlib import Path as _Path
    disk_path_str = _ud.normalize("NFC", str(disk_path))
    p = _Path(disk_path_str)
    if not p.exists():
        # APFS sometimes stores filenames in NFD even when the SQLite
        # row holds NFC. Look around the parent dir for an NFD twin
        # with the same trailing UUID prefix before giving up.
        match = None
        if p.parent.exists():
            head = p.name.split("-", 1)[0]
            for sibling in p.parent.iterdir():
                if sibling.name.startswith(head):
                    match = sibling
                    break
        if match is None:
            return {"inclusion": [], "exclusion": [], "schedule": [],
                    "protocol_summary": "",
                    "notes": [f"upload file missing on disk: {disk_path_str!r}"]}
        p = match

    try:
        head_bytes = p.open("rb").read(4)
    except OSError as exc:
        return {"inclusion": [], "exclusion": [], "schedule": [],
                "protocol_summary": "",
                "notes": [f"cannot read upload: {exc}"]}
    if head_bytes[:4] != b"PK\x03\x04":
        # .docx is a zip; anything else is either an old .doc (OLE,
        # starts with D0CF11E0), a renamed PDF (%PDF), or junk. Tell
        # the medic exactly why we can't parse it.
        if head_bytes[:4] == b"\xd0\xcf\x11\xe0":
            return {"inclusion": [], "exclusion": [], "schedule": [],
                    "protocol_summary": "",
                    "notes": ["uploaded file is a legacy .doc (OLE) — "
                              "please save as .docx in Word and re-upload"]}
        return {"inclusion": [], "exclusion": [], "schedule": [],
                "protocol_summary": "",
                "notes": [f"uploaded file is not a .docx (magic bytes "
                          f"{head_bytes!r}); maybe a renamed PDF or "
                          f"corrupted upload"]}

    try:
        # Hand python-docx an explicit binary handle so any zip-level
        # corruption surfaces as BadZipFile, not as the misleading
        # "Package not found" string.
        with p.open("rb") as f:
            doc = Document(f)
    except Exception as exc:  # noqa: BLE001
        return {"inclusion": [], "exclusion": [], "schedule": [],
                "protocol_summary": "",
                "notes": [f"docx parse failed: {exc}"]}

    # 3a. Try the LLM extractor first — it understands the whole doc
    # holistically (knows "剔除标准" == exclusion, knows "疗效评定" is
    # not a criterion, knows how to express age ranges as rule_dsl).
    doc_markdown = _docx_to_markdown(doc)
    llm_notes: list[str] = []
    try:
        llm_result = await _llm_extract_full(doc_markdown)
        if llm_result is not None:
            return {
                "study_title":       llm_result.get("study_title", ""),
                "short_code":        llm_result.get("short_code", ""),
                "phase":             llm_result.get("phase", ""),
                "primary_endpoint":  llm_result.get("primary_endpoint", ""),
                "inclusion":         llm_result.get("inclusion", []),
                "exclusion":         llm_result.get("exclusion", []),
                "schedule":          llm_result.get("schedule",  []),
                "protocol_summary":  llm_result.get("protocol_summary", ""),
                "notes":             llm_result.get("notes", []) + ["extracted by LLM"],
            }
        llm_notes.append("LLM returned no result — falling back to regex parser")
    except Exception as exc:  # noqa: BLE001
        logger.warning("protocol LLM extraction failed: %s — falling back to regex", exc)
        llm_notes.append(f"LLM extraction failed ({type(exc).__name__}); used regex fallback")

    # 3b. Regex fallback. Preserves the existing heuristic + classifier
    # so deployments without an LLM key (or with the gateway down) still
    # get a usable draft.
    sections = _split_into_sections(doc)
    inclusion_lines = sections.get("inclusion", [])
    exclusion_lines = sections.get("exclusion", [])
    schedule_lines  = sections.get("schedule", [])
    inclusion_draft = [_line_to_criterion(t, kind_hint="inclusion") for t in inclusion_lines]
    exclusion_draft = [_line_to_criterion(t, kind_hint="exclusion") for t in exclusion_lines]
    schedule_draft  = _lines_to_schedule(schedule_lines)
    summary = _short_summary(doc, sections)
    return {
        "inclusion": inclusion_draft,
        "exclusion": exclusion_draft,
        "schedule":  schedule_draft,
        "protocol_summary": summary,
        "notes": llm_notes,
    }


def _docx_to_markdown(doc) -> str:
    """Linearise paragraphs + tables into a single text blob for the LLM.

    Order matches the .docx flow (which preserves the protocol's logical
    section order), so the LLM can use positional cues — e.g. exclusion
    criteria typically follow inclusion, schedule typically follows
    treatment. Tables are joined cell-by-cell with `|` so the LLM still
    sees the row structure.
    """
    parts: list[str] = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


# Single-shot extraction prompt. The LLM sees the entire protocol text
# and returns one JSON object with the full criteria + schedule. The
# schema is identical to what the UI's batch-confirm page expects.
_PROTOCOL_EXTRACT_SYSTEM_PROMPT = """\
You are a structured-data extractor for clinical-trial protocols. Given
the full text of one trial protocol (a `.docx` linearised to plain
text), return ONE JSON object with:

  {
    "study_title":      "concise human-readable name (≤ 60 chars)",
    "short_code":       "3-12 char compact handle (uppercase + dashes)",
    "phase":            "I" | "II" | "III" | "IV" | "I/II" | "II/III" | "",
    "primary_endpoint": "short text or empty",
    "inclusion": [Criterion, ...],
    "exclusion": [Criterion, ...],
    "schedule":  [Visit, ...],
    "protocol_summary": "two-sentence plain-language summary",
    "notes": []
  }

`study_title` should be the actual trial name a medic would call this
study — e.g. "NSCLC IIIB cCRT + PD-1 维持 II 期" or "Hybrid RT NSCLC
Phase II". Do NOT just echo the docx filename. Do NOT include the
official long title verbatim if it's > 60 chars; produce a compact
human-readable form. If the protocol has an explicit name / short
title field, use that.

`short_code` is a stable engineering handle for this study — e.g.
"NEXUS-NSCLC-001", "HYBRID-RT-II", "ES-SCLC-EC". 3-12 chars, ASCII,
uppercase + dashes only. If the protocol declares its own code (like
"Protocol 编号: NEXUS-NSCLC-001"), use that verbatim.

`phase` is the trial phase as a string. Common values: "I", "I/II",
"II", "II/III", "III", "IV". If multi-phase (e.g. "Phase Ib/IIa")
collapse to the higher one's clean form ("I/II"). Empty string if the
protocol doesn't declare a phase.

Criterion = {
  "id":         "c-" + 8 hex chars (any unique id, you may make one up),
  "text":       single declarative sentence in the ORIGINAL language
                (Chinese stays Chinese, English stays English),
  "kind":       "auto-rule" | "auto-llm" | "manual",
  "rule_dsl":   string | null,
  "llm_prompt": string | null
}

Decide `kind` per criterion:
  - "auto-rule"   the criterion is a deterministic check against
                  structured patient facts (age, ECOG, AJCC stage,
                  pathology type, simple labs, driver mutation status).
                  Express the check as a `rule_dsl` string. Use these
                  patterns:
                     age BETWEEN 18 AND 75
                     ecog IN (0,1)
                     ajcc_stage CONTAINS 'IIIA' OR ajcc_stage CONTAINS 'IIIB'
                     driver_mutation = 'negative'
                     pathology = 'NSCLC'
                     anc >= 1.5    plt >= 100    hb >= 90
                     creat <= 1.5 * ULN    alt <= 2.5 * ULN
  - "auto-llm"    the criterion needs reading the patient's free-text
                  record to judge (e.g. "无活动性自身免疫病").
                  Set `llm_prompt` to a one-sentence instruction the
                  downstream LLM judge will use, in English.
  - "manual"      requires explicit medic action that can't be
                  automated (e.g. 签署知情同意书, MDT 评估为不可切除).

Visit = {
  "label":       visit name (e.g. "基线", "cCRT week 3"),
  "offset_days": days from D0=enrollment; negative for screening window,
  "assessments": list of assessment names done at this visit
}

Rules:
  - Extract ONLY actual eligibility criteria and visit points.
  - Do NOT include section headers, treatment regimen prose, statistical
    assumptions, stop-rules, CTCAE grading descriptions, consultation
    guidance, or table titles. Those are protocol body text, not
    criteria.
  - Each criterion gets one line of text. Do not concatenate criteria.
  - If the protocol uses "剔除标准" / "不予入组", treat them as exclusion.
  - Return STRICT JSON. No comments, no markdown fences, no trailing
    commas. If a field is unknown, use `null`; if a list is empty, use
    `[]`.
"""


async def _llm_extract_full(doc_markdown: str) -> Optional[dict]:
    """Call the LLM with the whole doc; return parsed JSON or None.

    Returning None (not raising) lets the caller decide between "use the
    LLM output" and "fall back to regex" cleanly. Raising signals an
    actual error (network, schema-validation failure) — the caller logs
    and still falls back.
    """
    if not doc_markdown.strip():
        return None
    try:
        from nexus_server import llm_gateway
    except Exception as exc:  # noqa: BLE001
        logger.info("llm_gateway unavailable: %s", exc)
        return None
    if not hasattr(llm_gateway, "call_llm"):
        return None

    # Cap the doc at ~40k chars to keep one round-trip cheap. Real
    # protocols rarely exceed that; if they do, the truncated tail is
    # almost always references / appendix and not material to eligibility.
    capped = doc_markdown[:40_000]
    truncated_note = (
        " (note: protocol text was truncated to 40k chars for the "
        "extractor — verify the schedule didn't lose late visits)"
        if len(doc_markdown) > 40_000 else ""
    )

    user_msg = (
        "Here is the protocol. Return the JSON object specified by your "
        "instructions, nothing else.\n\n"
        "------ PROTOCOL TEXT ------\n"
        + capped
    )
    content, _model, _stop, _tools = await llm_gateway.call_llm(
        messages=[{"role": "user", "content": user_msg}],
        system_prompt=_PROTOCOL_EXTRACT_SYSTEM_PROMPT,
        model=None,                      # use the gateway default
        temperature=0.1,                 # near-deterministic
        max_tokens=8192,
        tools=None,
    )
    if not content:
        return None

    parsed = _safe_parse_json(content)
    if parsed is None:
        logger.warning("LLM extract: failed to parse JSON, content head: %r",
                       content[:200])
        return None

    # Defensive normalisation: ensure every criterion has the keys the
    # UI's batch-confirm page reads. Drops obvious garbage (criteria
    # without `text`) silently rather than failing the import.
    def _norm_crits(items) -> list[dict]:
        out: list[dict] = []
        if not isinstance(items, list):
            return out
        for it in items:
            if not isinstance(it, dict):
                continue
            text = (it.get("text") or "").strip()
            if not text:
                continue
            out.append({
                "id":       it.get("id") or f"c-{abs(hash(text)) % 10**8:08x}",
                "text":     text,
                "kind":     it.get("kind") or "auto-llm",
                "rule_dsl": it.get("rule_dsl"),
                "llm_prompt":       it.get("llm_prompt"),
                "evidence_sources": it.get("evidence_sources"),
            })
        return out

    def _norm_visits(items) -> list[dict]:
        out: list[dict] = []
        if not isinstance(items, list):
            return out
        for it in items:
            if not isinstance(it, dict):
                continue
            label = (it.get("label") or "").strip()
            if not label:
                continue
            try:
                offset_days = int(it.get("offset_days") or 0)
            except (TypeError, ValueError):
                offset_days = 0
            assessments = it.get("assessments") or []
            if not isinstance(assessments, list):
                assessments = [str(assessments)]
            out.append({
                "label":       label,
                "offset_days": offset_days,
                "assessments": [str(a).strip() for a in assessments if a],
            })
        return out

    summary = (parsed.get("protocol_summary") or "").strip()
    if truncated_note:
        summary = (summary or "") + truncated_note

    # Sanitise study_title / short_code: keep these in a predictable
    # shape so the front-end's "New Study" form can drop them straight
    # into the display_name + short_code inputs without further
    # validation. A blank value just falls through to the filename-
    # derived default the caller already has.
    title = (parsed.get("study_title") or "").strip()
    if len(title) > 80:                          # hard ceiling
        title = title[:77].rstrip() + "…"
    code = (parsed.get("short_code") or "").strip()
    # Strip everything but ASCII letters / digits / dash / underscore.
    code = re.sub(r"[^A-Za-z0-9_-]", "", code)[:16].upper()
    phase = (parsed.get("phase") or "").strip()
    primary = (parsed.get("primary_endpoint") or "").strip()

    return {
        "study_title":      title,
        "short_code":       code,
        "phase":            phase,
        "primary_endpoint": primary,
        "inclusion":        _norm_crits(parsed.get("inclusion")),
        "exclusion":        _norm_crits(parsed.get("exclusion")),
        "schedule":         _norm_visits(parsed.get("schedule")),
        "protocol_summary": summary,
        "notes":            parsed.get("notes") or [],
    }


def _safe_parse_json(content: str) -> Optional[dict]:
    """Parse JSON from LLM output. Tolerates ```json fences and stray
    leading/trailing prose. Returns None on persistent failure rather
    than raising — the caller already has a regex-based fallback.
    """
    content = content.strip()
    # Strip code fences if the LLM ignored "no markdown" and wrapped the
    # JSON in ```json ... ```.
    if content.startswith("```"):
        content = re.sub(r"^```[a-zA-Z]*\n?", "", content)
        content = re.sub(r"\n?```\s*$", "", content)
    # First try as-is.
    try:
        return json.loads(content)
    except json.JSONDecodeError:
        pass
    # Find the largest balanced {...} substring as a fallback.
    start = content.find("{")
    end   = content.rfind("}")
    if start >= 0 and end > start:
        candidate = content[start:end + 1]
        try:
            return json.loads(candidate)
        except json.JSONDecodeError:
            return None
    return None


# ─────────────────────────────────────────────────────────────────────
# Section splitting
# ─────────────────────────────────────────────────────────────────────


def _split_into_sections(doc) -> dict[str, list[str]]:
    """Walk paragraphs + tables, group lines by section header.

    State machine:
        current = None                — nothing being collected
        current = inclusion/exclusion/schedule — collecting that section
        encountering an end-of-criteria header (治疗方案 / 终点 / ...) →
          back to None so subsequent body text doesn't leak into the
          last-active criteria bucket.

    A per-section cap (50) protects against runaway harvesting if a
    protocol genuinely has thousands of "criteria-shaped" lines.
    """
    out: dict[str, list[str]] = {"inclusion": [], "exclusion": [], "schedule": []}
    current: Optional[str] = None

    def _switch(new_current: Optional[str]) -> None:
        nonlocal current
        current = new_current

    def _maybe_collect(text: str) -> None:
        if not current:
            return
        if len(out[current]) >= 50:
            return
        if not _looks_like_criterion(text):
            return
        out[current].append(text)

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        # Order matters: criteria headers should beat the end-of-section
        # detector (e.g. "5. 安全性入选标准" should open inclusion, not
        # close it via the "安全性" match).
        if _INCL_HEADER.search(text):
            _switch("inclusion"); continue
        if _EXCL_HEADER.search(text):
            _switch("exclusion"); continue
        if _SCHEDULE_HEADER.search(text):
            _switch("schedule"); continue
        if _END_OF_CRITERIA_HEADER.search(text):
            _switch(None); continue
        _maybe_collect(text)

    # Tables: many CN protocols put inclusion/exclusion in tables.
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text]
            joined = " | ".join(cells)
            if not joined:
                continue
            if _INCL_HEADER.search(joined):
                _switch("inclusion"); continue
            if _EXCL_HEADER.search(joined):
                _switch("exclusion"); continue
            if _SCHEDULE_HEADER.search(joined):
                _switch("schedule"); continue
            if _END_OF_CRITERIA_HEADER.search(joined):
                _switch(None); continue
            _maybe_collect(joined)
    return out


_BULLET_RE = re.compile(r"^([a-zA-Z]|[0-9]+|\([0-9]+\)|[一二三四五六七八九十]+)[\.、\)）]")

# Phrases that strongly indicate a paragraph is NOT a criterion even
# when it incidentally contains a comparator. CTCAE descriptions,
# response-evaluation prose, and treatment-action sentences all match
# the old keyword heuristic but aren't eligibility-actionable.
_NOT_A_CRITERION_PATTERNS = (
    re.compile(r"^描述\s*[|｜]"),                    # CTCAE grading prose
    re.compile(r"^级别\s*[|｜]"),                    # grading scale rows
    re.compile(r"^(疗效|有效性)评[价定]"),
    re.compile(r"^检查及对症治疗"),
    re.compile(r"^(会诊|转诊|医生指示)"),
    re.compile(r"^(给药|剂量|滴注|静滴)"),
    re.compile(r"^(I+|IV)\s*级\s*推荐"),             # NCCN guideline lines
    re.compile(r"^(\d+\s*[、.]\s*)?[未无]\s*症状"),  # "无症状" CTCAE entries
)


def _looks_like_criterion(text: str) -> bool:
    """Return True only when the paragraph plausibly describes a single
    enumerable eligibility/exclusion criterion. False-positive rate
    matters more than false-negative here: any line we wrongly admit
    ends up in the medic's review screen as junk; missing a real
    criterion is recoverable via the "+ Add" button.
    """
    if len(text) < 4:
        return False
    # Drop anything looking like protocol body text — CTCAE grading
    # rows, treatment / consultation guidance, etc.
    for pat in _NOT_A_CRITERION_PATTERNS:
        if pat.search(text):
            return False
    # Bullet-prefixed lines (1./二、/(3)/...) are the strongest signal.
    if _BULLET_RE.match(text):
        return True
    # Long declarative sentences with a *structural* eligibility marker
    # — age comparator, ECOG, stage, lab cutoff. The old single-char
    # keywords ("无"/"有"/"未") were too permissive; CTCAE descriptions
    # like "无症状" matched them.
    structural = (
        "年龄", "ECOG", "ECOG PS",
        "≥", "≤", "BETWEEN", "IN (",
        "AJCC", "RECIST",
    )
    if any(k in text for k in structural):
        return len(text) >= 10
    return False


# ─────────────────────────────────────────────────────────────────────
# Heuristic → CriterionDef draft
# ─────────────────────────────────────────────────────────────────────


_AGE_RANGE = re.compile(r"(\d{1,2})\s*[-–~至]\s*(\d{1,2})\s*岁")
_AGE_GTE   = re.compile(r"年龄[≥>]=?\s*(\d{1,2})")
_ECOG      = re.compile(r"ECOG\s*(?:PS)?\s*(?:评分)?\s*([0-9])\s*[-–~至到]\s*([0-9])")
_DRIVER_NEG = re.compile(r"驱动(?:基因)?(?:全)?阴性")
_DRIVER_POS = re.compile(r"驱动(?:基因)?(?:阳性|有突变)")
_STAGE     = re.compile(r"(IV[ABC]?|III[ABC]?|II[AB]?|I[A]?)\s*期")
_CONSENT   = re.compile(r"知情同意")
_NSCLC     = re.compile(r"NSCLC|非小细胞肺癌")
_SCLC      = re.compile(r"SCLC|小细胞肺癌")


def _line_to_criterion(text: str, *, kind_hint: str) -> dict:
    cid = f"c-{abs(hash(text)) % 10**8:08x}"
    base = {"id": cid, "text": text, "kind": "manual", "rule_dsl": None,
            "llm_prompt": None, "evidence_sources": None}

    # Age range
    m = _AGE_RANGE.search(text)
    if m:
        lo, hi = int(m.group(1)), int(m.group(2))
        base.update(kind="auto-rule", rule_dsl=f"age BETWEEN {lo} AND {hi}")
        return base

    m = _AGE_GTE.search(text)
    if m:
        lo = int(m.group(1))
        base.update(kind="auto-rule", rule_dsl=f"age >= {lo}")
        return base

    m = _ECOG.search(text)
    if m:
        lo, hi = int(m.group(1)), int(m.group(2))
        vals = ",".join(str(i) for i in range(lo, hi + 1))
        base.update(kind="auto-rule", rule_dsl=f"ecog IN ({vals})")
        return base

    if _DRIVER_NEG.search(text):
        base.update(kind="auto-rule", rule_dsl="driver_mutation = 'negative'")
        return base
    if _DRIVER_POS.search(text):
        base.update(kind="auto-rule", rule_dsl="driver_mutation = 'positive'")
        return base

    m = _STAGE.search(text)
    if m:
        base.update(kind="auto-rule",
                    rule_dsl=f"ajcc_stage CONTAINS '{m.group(1)}'")
        return base

    if _CONSENT.search(text):
        base.update(kind="manual")
        return base

    if _NSCLC.search(text):
        base.update(kind="auto-rule",
                    rule_dsl="pathology IN ('adenocarcinoma','SCC','adenosquamous','LCC','NSCLC')")
        return base
    if _SCLC.search(text):
        base.update(kind="auto-rule", rule_dsl="pathology = 'SCLC'")
        return base

    # Default: leave as 'auto-llm' so the LLM judge picks it up at
    # eligibility time. Free-text criteria like "病人状态适合放疗"
    # are the canonical auto-llm case.
    base.update(kind="auto-llm",
                llm_prompt="Evaluate whether this criterion is met based on "
                           "the structured patient_facts + free-text notes.",
                evidence_sources=["soap", "pathology", "imaging"])
    return base


# ─────────────────────────────────────────────────────────────────────
# Schedule extraction (very heuristic; LLM refines)
# ─────────────────────────────────────────────────────────────────────


_VISIT_OFFSET_DAYS = {
    "基线": 0, "Baseline": 0, "Day 1": 1,
    "1 周": 7, "1周": 7, "Week 1": 7,
    "2 周": 14, "2周": 14, "Week 2": 14,
    "4 周": 28, "4周": 28, "Week 4": 28,
    "1 月": 30, "1月": 30, "1个月": 30,
    "3 月": 90, "3月": 90, "3个月": 90, "Q3M": 90,
    "6 月": 180, "6月": 180, "6个月": 180, "Q6M": 180,
    "1 年": 365, "1年": 365, "Year 1": 365,
}


def _lines_to_schedule(lines: list[str]) -> list[dict]:
    out: list[dict] = []
    for ln in lines:
        for marker, off in _VISIT_OFFSET_DAYS.items():
            if marker in ln:
                # crude: pull keywords like 'CT'/'MRI'/'抽血' as assessments
                kinds = []
                for kw in ("CT", "MRI", "PET", "骨扫描", "抽血",
                           "血常规", "心电图", "TFT", "肿瘤标志物"):
                    if kw in ln:
                        kinds.append(_kw_to_kind(kw))
                if not kinds:
                    kinds = ["unspecified"]
                out.append({
                    "label": marker, "offset_days": off,
                    "assessments": list(dict.fromkeys(kinds)),  # dedup, preserve order
                })
                break
    return out


def _kw_to_kind(kw: str) -> str:
    m = {
        "CT": "imaging_ct", "MRI": "imaging_mri", "PET": "imaging_pet_ct",
        "骨扫描": "imaging_bone_scan", "抽血": "lab_panel",
        "血常规": "cbc", "心电图": "ecg", "TFT": "tft",
        "肿瘤标志物": "tumor_markers",
    }
    return m.get(kw, "unspecified")


# ─────────────────────────────────────────────────────────────────────
# Optional LLM refinement
# ─────────────────────────────────────────────────────────────────────


def _llm_refine(drafts: list[dict], side: str) -> list[dict]:
    """Best-effort LLM pass to add llm_prompt + tighten rule_dsl. Only
    triggered when an LLM gateway is configured."""
    return drafts  # No-op for now; the heuristic draft is already
                   # solid for the three reference protocols. The LLM
                   # pass adds value mainly on novel protocols and can
                   # be wired up Phase 2.5 without changing the
                   # protocol_parser contract.


def _short_summary(doc, sections: dict[str, list[str]]) -> str:
    """First 2-3 paragraphs of the doc, no LLM."""
    paras = [p.text.strip() for p in doc.paragraphs if (p.text or "").strip()][:5]
    return "\n".join(paras)[:600]
