"""Writing Studio HTTP surface (P1 MVP).

See docs/design/WRITING_STUDIO_DESIGN.docx. Implements the server side
of the in-app writing mode: documents + version snapshots + data
reference chips (auto de-identified at insert time) + selection polish
(SSE) + PHI scan + docx export gated on PHI resolution.

Endpoints (all mounted under /api/v1, all auth-gated via
``Depends(get_current_user)``; every query is user_id-scoped so one
medic can never touch another's documents):

  GET    /docs                              list (id, title, updated_at, ref_count)
  POST   /docs                              create {title}
  GET    /docs/{id}                         full doc + references
  PUT    /docs/{id}                         save {title?, body?} (+snapshot on body change)
  DELETE /docs/{id}                         delete doc + refs + snapshots
  GET    /docs/{id}/snapshots               version chain (latest 50)
  POST   /docs/{id}/snapshots/{sid}/restore restore a snapshot
  POST   /docs/{id}/references              insert a de-identified data chip
  POST   /docs/{id}/polish                  SSE selection rewrite
  GET    /docs/{id}/chat                    co-writing chat history
  POST   /docs/{id}/chat                    SSE co-writing turn (may
                                            regenerate the doc body)
  POST   /docs/{id}/phi-scan                regex + roster-name PHI findings
  POST   /docs/{id}/export                  expand chips → .docx (PHI gate)

De-identification rules (design §5):
  patient  — name/initials → 'P-' + patient_hash[:6]; birth data → age;
             absolute dates in the timeline → relative 'D+N周' (falls
             back to month precision YYYY-MM when unparseable); MRN
             scrubbed. Snapshot is FROZEN at insert time.
  study    — aggregate counts only ('42/60 例'); roster members appear
             as screening codes (SHORTCODE-007), never hashes/names.
  file     — reuses the distill/summary text stored at upload time.

Audit: every reference insert emits a DOC_REFERENCE_CREATED event into
twin_event_log via Store.emit_and_apply (who / when / which doc /
which target / which granularity).
"""

from __future__ import annotations

import json
import logging
import re
import sqlite3
import uuid
from datetime import date, datetime, timezone
from io import BytesIO
from typing import Any, AsyncIterator, Optional

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import JSONResponse, Response, StreamingResponse
from pydantic import BaseModel, Field

from nexus_server.auth.routes import get_current_user
from nexus_server.database import get_db_connection

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/v1", tags=["writing"])

# Version-snapshot retention per doc. Older rows are pruned on save.
_SNAPSHOT_CAP = 50

# {{ref:UUID}} chip placeholder in doc bodies.
_REF_PLACEHOLDER_RE = re.compile(r"\{\{ref:([0-9a-fA-F-]{8,64})\}\}")

# Numbers for the provenance (hallucinated-value) check.
_NUM_RE = re.compile(r"\d+(?:\.\d+)?%?")

# ── PHI regexes (layer 1) ────────────────────────────────────────────
# Full dates: 2025年3月14日 / 2025-03-14 / 2025/3/14 / 2025.3.14.
_PHI_DATE_RE = re.compile(
    r"((?:19|20)\d{2})\s*[年/\-.]\s*(\d{1,2})\s*[月/\-.]\s*(\d{1,2})\s*日?"
)
# 18-digit CN resident ID (last char may be X).
_PHI_ID_RE = re.compile(r"(?<![\dXx])\d{17}[\dXx](?![\dXx\d])")
# CN mobile numbers.
_PHI_PHONE_RE = re.compile(r"(?<!\d)1[3-9]\d{9}(?!\d)")


# ─────────────────────────────────────────────────────────────────────
# Schema (defensive mirror of database.init_db — idempotent)
# ─────────────────────────────────────────────────────────────────────

def _ensure_schema(conn: sqlite3.Connection) -> None:
    """Bring up the Writing Studio tables if a stale deployment booted
    before database.init_db learned about them. Idempotent."""
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS docs (
            id         TEXT PRIMARY KEY,
            user_id    TEXT NOT NULL,
            title      TEXT NOT NULL DEFAULT '',
            body       TEXT NOT NULL DEFAULT '',
            created_at TIMESTAMP NOT NULL,
            updated_at TIMESTAMP NOT NULL
        )
        """
    )
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS doc_references (
            id           TEXT PRIMARY KEY,
            doc_id       TEXT NOT NULL,
            user_id      TEXT NOT NULL,
            ref_type     TEXT NOT NULL,
            target_id    TEXT NOT NULL,
            granularity  TEXT NOT NULL DEFAULT '',
            snapshot     TEXT NOT NULL DEFAULT '',
            source_nodes TEXT NOT NULL DEFAULT '[]',
            created_at   TIMESTAMP NOT NULL
        )
        """
    )
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS doc_snapshots (
            id         INTEGER PRIMARY KEY AUTOINCREMENT,
            doc_id     TEXT NOT NULL,
            user_id    TEXT NOT NULL,
            body       TEXT NOT NULL,
            label      TEXT NOT NULL DEFAULT '',
            created_at TIMESTAMP NOT NULL
        )
        """
    )
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS doc_chat_messages (
            id          TEXT PRIMARY KEY,
            doc_id      TEXT NOT NULL,
            user_id     TEXT NOT NULL,
            role        TEXT NOT NULL,
            text        TEXT NOT NULL DEFAULT '',
            doc_applied INTEGER NOT NULL DEFAULT 0,
            created_at  TIMESTAMP NOT NULL
        )
        """
    )


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


def _get_doc_or_404(
    conn: sqlite3.Connection, user_id: str, doc_id: str,
) -> sqlite3.Row:
    row = conn.execute(
        "SELECT id, user_id, title, body, created_at, updated_at "
        "FROM docs WHERE user_id = ? AND id = ?",
        (user_id, doc_id),
    ).fetchone()
    if row is None:
        raise HTTPException(status_code=404, detail="doc not found")
    return row


def _sse(event: dict) -> str:
    """Serialise a chunk as an SSE message (same convention as
    chat_router)."""
    return f"data: {json.dumps(event, ensure_ascii=False)}\n\n"


# ─────────────────────────────────────────────────────────────────────
# Models
# ─────────────────────────────────────────────────────────────────────

class DocCreateRequest(BaseModel):
    title: str = Field("", max_length=500)


class DocUpdateRequest(BaseModel):
    title: Optional[str] = Field(None, max_length=500)
    body: Optional[str] = None
    snapshot_label: str = Field("save", max_length=200)


class ReferenceCreateRequest(BaseModel):
    ref_type: str = Field(..., pattern="^(patient|study|file)$")
    target_id: str = Field(..., min_length=1, max_length=256)
    # patient: basics | timeline    study: progress | roster
    # file: summary (default)
    granularity: str = Field("basics", max_length=64)


class PolishRequest(BaseModel):
    selection: str = Field(..., min_length=1)
    instruction: str = ""
    ref_ids: list[str] = []


class ChatRequest(BaseModel):
    message: str = Field(..., min_length=1)
    # Optional subset of doc reference ids to inject as context; when
    # empty/omitted ALL of the doc's reference snapshots are injected.
    ref_ids: list[str] = []
    # Explicit per-message skill invocation (the "/" menu in the
    # composer). Each name must be an installed + enabled skill for
    # this user — unknown / disabled names are silently dropped.
    # Independent of this list, skills flagged auto_apply=1 in
    # user_skill_prefs are injected on EVERY turn. Same contract as
    # chat_router.ChatRequest; see skills_router.build_skills_block.
    skills: list[str] = []
    # File attachments staged in the writing composer (📎 button /
    # paste / drag-drop). Each entry is {file_id, name}. file_id must
    # reference an ``uploads`` row owned by the current user; unknown /
    # foreign ids are silently skipped (same contract as the v2 chat
    # path in chat_router). Extracted / distilled text is injected into
    # the co-writing system prompt as '## 用户附件: {name}' sections —
    # after the reference snapshots, before the skills block.
    attachments: list[dict] = []


class PhiScanRequest(BaseModel):
    # Optional override so the frontend can scan an unsaved draft.
    body: Optional[str] = None


class ExportRequest(BaseModel):
    # Each resolution addresses one phi-scan finding. Matching is by
    # exact (start, end) span OR by excerpt string. Optional
    # ``replacement`` rewrites the flagged text before export;
    # a resolution without replacement means "explicitly keep".
    resolutions: list[dict] = []
    include_sources: bool = False


# ─────────────────────────────────────────────────────────────────────
# De-identification helpers
# ─────────────────────────────────────────────────────────────────────

_GRANULARITY_LABELS = {
    "basics":   "基本特征",
    "timeline": "治疗时间线",
    "progress": "研究进展",
    "roster":   "roster 快照",
    "summary":  "摘要",
}


def _patient_code(patient_hash: str) -> str:
    return "P-" + (patient_hash or "")[:6]


def _scrub_terms(text: str, terms: list[str], replacement: str) -> str:
    """Replace every occurrence of each identifying term (and its
    whitespace-free variant) with ``replacement``."""
    out = text
    variants: list[str] = []
    for t in terms:
        t = (t or "").strip()
        if len(t) < 2:
            continue
        variants.append(t)
        squeezed = re.sub(r"\s+", "", t)
        if squeezed != t and len(squeezed) >= 2:
            variants.append(squeezed)
    # Longest first so "张三丰" wins over "张三".
    for v in sorted(set(variants), key=len, reverse=True):
        out = out.replace(v, replacement)
    return out


def _parse_date(y: str, m: str, d: str) -> Optional[date]:
    try:
        return date(int(y), int(m), int(d))
    except ValueError:
        return None


def _relativize_dates(text: str, anchor: Optional[date]) -> str:
    """Rewrite absolute full dates to 'D+N周' relative form (anchored
    at the earliest date in the reference material). Unparseable dates
    degrade to month precision YYYY-MM."""

    def _sub(m: re.Match) -> str:
        d = _parse_date(m.group(1), m.group(2), m.group(3))
        if d is None:
            return f"{m.group(1)}-{int(m.group(2) or 1):02d}"
        if anchor is None:
            return f"{d.year}-{d.month:02d}"
        delta_days = (d - anchor).days
        if delta_days == 0:
            return "D0"
        weeks = delta_days // 7 if delta_days > 0 else -((-delta_days) // 7)
        sign = "+" if delta_days > 0 else "-"
        return f"D{sign}{abs(weeks)}周"

    return _PHI_DATE_RE.sub(_sub, text)


def _earliest_date(texts: list[str]) -> Optional[date]:
    found: list[date] = []
    for t in texts:
        for m in _PHI_DATE_RE.finditer(t or ""):
            d = _parse_date(m.group(1), m.group(2), m.group(3))
            if d is not None:
                found.append(d)
    return min(found) if found else None


def _node_display(content: dict) -> str:
    """Best-effort human string for one clinical_graph node payload."""
    for key in ("label", "name", "text", "summary", "value"):
        v = content.get(key)
        if isinstance(v, str) and v.strip():
            return v.strip()
    try:
        return json.dumps(content, ensure_ascii=False)[:120]
    except Exception:  # noqa: BLE001
        return ""


def _build_patient_snapshot(
    conn: sqlite3.Connection, user_id: str, patient_hash: str,
    granularity: str,
) -> tuple[str, list[Any], str]:
    """Compose the de-identified patient snapshot.

    Returns (snapshot_text, source_node_ids, chip_label). Defensive by
    design: any missing table / row just shortens the snapshot — this
    function never raises for missing data.
    """
    code = _patient_code(patient_hash)
    scrub: list[str] = []
    lines: list[str] = []
    source_nodes: list[Any] = []

    prow = None
    try:
        prow = conn.execute(
            "SELECT initials, mrn, age_value, age_group, sex, "
            "       chief_complaint, notes "
            "FROM patients WHERE user_id = ? AND patient_hash = ?",
            (user_id, patient_hash),
        ).fetchone()
    except sqlite3.Error as e:
        logger.debug("patients lookup failed: %s", e)

    if prow is not None:
        initials, mrn = str(prow[0] or ""), str(prow[1] or "")
        age_value, age_group = int(prow[2] or 0), str(prow[3] or "")
        sex, chief = str(prow[4] or ""), str(prow[5] or "")
        scrub = [initials, mrn]
        sex_label = {"M": "男", "F": "女"}.get(sex, "")
        age_label = (
            f"{age_value}岁" if age_value > 0
            else (f"{age_group}岁段" if age_group else "")
        )
        header = f"患者 {code}"
        details = "，".join(x for x in (age_label, sex_label) if x)
        if details:
            header += f"（{details}）"
        lines.append(header)
        if chief:
            lines.append(f"主诉：{chief}")
    else:
        lines.append(f"患者 {code}")

    if granularity == "timeline":
        rows: list[sqlite3.Row] = []
        try:
            rows = conn.execute(
                "SELECT node_id, node_type, content_json, encounter_id, "
                "       updated_at "
                "FROM clinical_graph_nodes "
                "WHERE user_id = ? AND patient_hash = ? "
                "  AND node_type IN "
                "      ('finding','measurement','med','study','encounter') "
                "ORDER BY updated_at ASC LIMIT 100",
                (user_id, patient_hash),
            ).fetchall()
        except sqlite3.Error as e:
            logger.debug("timeline nodes lookup failed: %s", e)

        entries: list[tuple[str, str]] = []  # (type_label, text)
        type_labels = {
            "finding": "发现", "measurement": "测量", "med": "用药",
            "study": "检查", "encounter": "就诊",
        }
        raw_texts: list[str] = []
        for r in rows:
            try:
                content = json.loads(r[2] or "{}")
            except (TypeError, ValueError):
                content = {}
            display = _node_display(content)
            if not display:
                continue
            date_hint = str(
                content.get("date") or content.get("study_date") or ""
            )
            text = f"{display}" + (f"（{date_hint}）" if date_hint else "")
            raw_texts.append(text)
            entries.append((type_labels.get(str(r[1]), str(r[1])), text))
            source_nodes.append(r[0])

        anchor = _earliest_date(raw_texts)
        if entries:
            lines.append("治疗时间线：")
            for tl, text in entries:
                lines.append(f"- [{tl}] {_relativize_dates(text, anchor)}")
        else:
            lines.append("治疗时间线：暂无结构化记录")

    snapshot = "\n".join(lines)
    # Final de-identification pass: names/MRN → code, then any
    # residual full dates / IDs / phone numbers scrubbed.
    snapshot = _scrub_terms(snapshot, scrub, code)
    snapshot = _relativize_dates(snapshot, None) if granularity != "timeline" else snapshot
    snapshot = _PHI_ID_RE.sub("[已脱敏]", snapshot)
    snapshot = _PHI_PHONE_RE.sub("[已脱敏]", snapshot)

    label = f"{code}·{_GRANULARITY_LABELS.get(granularity, granularity)}"
    return snapshot, source_nodes, label


def _build_study_snapshot(
    conn: sqlite3.Connection, user_id: str, study_id: str, granularity: str,
) -> tuple[str, list[Any], str]:
    """Study reference snapshot — aggregate counts / roster codes only,
    no per-patient identifiers (design §5: 成员患者一律以编号呈现)."""
    srow = None
    try:
        srow = conn.execute(
            "SELECT display_name, short_code, target_n, status "
            "FROM research_studies WHERE user_id = ? AND study_id = ?",
            (user_id, study_id),
        ).fetchone()
    except sqlite3.Error as e:
        logger.debug("research_studies lookup failed: %s", e)
    if srow is None:
        raise HTTPException(status_code=404, detail="study not found")

    display_name = str(srow[0] or "")
    short_code = str(srow[1] or "") or study_id[:8]
    target_n = srow[2]
    status = str(srow[3] or "")

    lines: list[str] = [f"研究 {short_code}（{display_name}）"]
    source_nodes: list[Any] = []

    def _count(sql: str, params: tuple) -> int:
        try:
            row = conn.execute(sql, params).fetchone()
            return int(row[0] or 0) if row else 0
        except sqlite3.Error:
            return 0

    enrolled = _count(
        "SELECT COUNT(*) FROM study_enrollments "
        "WHERE user_id = ? AND study_id = ? "
        "  AND status IN ('enrolled','completed')",
        (user_id, study_id),
    )
    screened = _count(
        "SELECT COUNT(*) FROM screening_evaluations "
        "WHERE user_id = ? AND study_id = ?",
        (user_id, study_id),
    )

    if granularity == "roster":
        lines.append(f"roster 快照（入组 {enrolled} 例）：")
        try:
            rows = conn.execute(
                "SELECT enrollment_seq, status, arm "
                "FROM study_enrollments "
                "WHERE user_id = ? AND study_id = ? "
                "ORDER BY enrollment_seq ASC LIMIT 200",
                (user_id, study_id),
            ).fetchall()
        except sqlite3.Error:
            rows = []
        for r in rows:
            seq, st, arm = int(r[0] or 0), str(r[1] or ""), str(r[2] or "")
            line = f"- {short_code}-{seq:03d}：{st}"
            if arm:
                line += f"（{arm} 组）"
            lines.append(line)
        if not rows:
            lines.append("- 暂无入组记录")
    else:  # 'progress' (default)
        target = str(target_n) if target_n else "?"
        lines.append(f"入组进度：{enrolled}/{target} 例")
        lines.append(f"筛查评估：{screened} 人次")
        if status:
            lines.append(f"研究状态：{status}")

    label = (
        f"{short_code}·{_GRANULARITY_LABELS.get(granularity, granularity)}"
    )
    return "\n".join(lines), source_nodes, label


def _build_file_snapshot(
    conn: sqlite3.Connection, user_id: str, file_id: str, granularity: str,
) -> tuple[str, list[Any], str]:
    """File/literature reference — reuse the distill/summary text
    captured at upload time (quick_scan_summary / memory_summary /
    extracted_text head)."""
    urow = None
    try:
        urow = conn.execute(
            "SELECT name, extracted_text, "
            "       COALESCE(quick_scan_summary, ''), "
            "       COALESCE(memory_summary, '') "
            "FROM uploads WHERE user_id = ? AND file_id = ?",
            (user_id, file_id),
        ).fetchone()
    except sqlite3.Error:
        # Older uploads schema without the summary columns.
        try:
            r = conn.execute(
                "SELECT name, extracted_text FROM uploads "
                "WHERE user_id = ? AND file_id = ?",
                (user_id, file_id),
            ).fetchone()
            urow = (r[0], r[1], "", "") if r else None
        except sqlite3.Error as e:
            logger.debug("uploads lookup failed: %s", e)
            urow = None
    if urow is None:
        raise HTTPException(status_code=404, detail="file not found")

    name = str(urow[0] or file_id)
    extracted = str(urow[1] or "").strip()
    qs_summary = str(urow[2] or "").strip()
    mem_summary = str(urow[3] or "").strip()

    summary = qs_summary or mem_summary or extracted[:1200]
    lines = [f"文件《{name}》"]
    if summary:
        lines.append(summary)
    else:
        lines.append("（无可用摘要）")
    return "\n".join(lines), [file_id], name


def _build_reference_snapshot(
    conn: sqlite3.Connection, user_id: str, ref_type: str, target_id: str,
    granularity: str,
) -> tuple[str, list[Any], str]:
    if ref_type == "patient":
        return _build_patient_snapshot(conn, user_id, target_id, granularity)
    if ref_type == "study":
        return _build_study_snapshot(conn, user_id, target_id, granularity)
    return _build_file_snapshot(conn, user_id, target_id, granularity)


def _chip_label_for_row(
    conn: sqlite3.Connection, user_id: str, ref_row: sqlite3.Row,
) -> str:
    """Recompute a chip label from a doc_references row (cheap lookups,
    all defensive)."""
    ref_type = str(ref_row["ref_type"])
    target_id = str(ref_row["target_id"])
    granularity = str(ref_row["granularity"])
    g_label = _GRANULARITY_LABELS.get(granularity, granularity)
    if ref_type == "patient":
        return f"{_patient_code(target_id)}·{g_label}"
    if ref_type == "study":
        short = target_id[:8]
        try:
            r = conn.execute(
                "SELECT short_code FROM research_studies "
                "WHERE user_id = ? AND study_id = ?",
                (user_id, target_id),
            ).fetchone()
            if r and r[0]:
                short = str(r[0])
        except sqlite3.Error:
            pass
        return f"{short}·{g_label}"
    # file
    try:
        r = conn.execute(
            "SELECT name FROM uploads WHERE user_id = ? AND file_id = ?",
            (user_id, target_id),
        ).fetchone()
        if r and r[0]:
            return str(r[0])
    except sqlite3.Error:
        pass
    return target_id[:16]


# ─────────────────────────────────────────────────────────────────────
# PHI scan
# ─────────────────────────────────────────────────────────────────────

def _placeholder_spans(body: str) -> list[tuple[int, int]]:
    return [(m.start(), m.end()) for m in _REF_PLACEHOLDER_RE.finditer(body)]


def _inside_any(start: int, end: int, spans: list[tuple[int, int]]) -> bool:
    return any(s <= start and end <= e for s, e in spans)


def _phi_scan_body(
    conn: sqlite3.Connection, user_id: str, body: str,
) -> list[dict]:
    """Two-layer PHI scan over manually-typed doc text.

    Layer 1 — regex: full dates, 18-digit resident IDs, CN mobile
    numbers. Layer 2 — high-precision roster match: THIS user's
    patients table names (initials field + whitespace-free variants)
    flagged wherever they appear.

    Chip placeholders ({{ref:...}}) are skipped — their content was
    de-identified at insert time.
    """
    findings: list[dict] = []
    skip_spans = _placeholder_spans(body)
    taken: list[tuple[int, int]] = []

    def _add(kind: str, m_start: int, m_end: int, excerpt: str,
             suggestion: str) -> None:
        if _inside_any(m_start, m_end, skip_spans):
            return
        # Overlap guard: an 18-digit ID also matches the phone regex's
        # digit run etc. First (higher-priority) finding wins the span.
        if any(not (m_end <= s or m_start >= e) for s, e in taken):
            return
        taken.append((m_start, m_end))
        findings.append({
            "kind": kind, "excerpt": excerpt,
            "start": m_start, "end": m_end,
            "suggestion": suggestion,
        })

    # Layer 2 first — patient names are the highest-value catch and
    # should win any span overlap with layer-1 patterns.
    try:
        rows = conn.execute(
            "SELECT patient_hash, initials FROM patients "
            "WHERE user_id = ? AND initials != ''",
            (user_id,),
        ).fetchall()
    except sqlite3.Error:
        rows = []
    for r in rows:
        phash, name = str(r[0]), str(r[1] or "").strip()
        if len(name) < 2:
            continue  # single chars are too noisy to flag
        code = _patient_code(phash)
        variants = {name, re.sub(r"\s+", "", name)}
        for v in variants:
            if len(v) < 2:
                continue
            for m in re.finditer(re.escape(v), body):
                _add(
                    "patient_name", m.start(), m.end(), m.group(0),
                    f"替换为研究编号 {code}",
                )

    # Layer 1 — IDs before phones (an ID contains a phone-shaped run).
    for m in _PHI_ID_RE.finditer(body):
        _add("id_number", m.start(), m.end(), m.group(0),
             "移除证件号或替换为研究编号")
    for m in _PHI_PHONE_RE.finditer(body):
        _add("phone", m.start(), m.end(), m.group(0), "移除电话号码")
    for m in _PHI_DATE_RE.finditer(body):
        _add("exact_date", m.start(), m.end(), m.group(0), "改为相对时间")

    findings.sort(key=lambda f: f["start"])
    return findings


# ─────────────────────────────────────────────────────────────────────
# Docs CRUD
# ─────────────────────────────────────────────────────────────────────

@router.get("/docs")
async def list_docs(
    current_user: str = Depends(get_current_user),
) -> dict[str, Any]:
    with get_db_connection() as conn:
        _ensure_schema(conn)
        rows = conn.execute(
            "SELECT d.id, d.title, d.updated_at, "
            "       (SELECT COUNT(*) FROM doc_references r "
            "         WHERE r.user_id = d.user_id AND r.doc_id = d.id) "
            "       AS ref_count "
            "FROM docs d WHERE d.user_id = ? "
            "ORDER BY d.updated_at DESC",
            (current_user,),
        ).fetchall()
        return {
            "docs": [
                {
                    "id": r[0], "title": r[1],
                    "updated_at": r[2], "ref_count": int(r[3] or 0),
                }
                for r in rows
            ]
        }


@router.post("/docs")
async def create_doc(
    req: DocCreateRequest,
    current_user: str = Depends(get_current_user),
) -> dict[str, Any]:
    doc_id = str(uuid.uuid4())
    now = _now()
    with get_db_connection() as conn:
        _ensure_schema(conn)
        conn.execute(
            "INSERT INTO docs (id, user_id, title, body, created_at, "
            "updated_at) VALUES (?, ?, ?, '', ?, ?)",
            (doc_id, current_user, req.title, now, now),
        )
        conn.commit()
    return {
        "id": doc_id, "title": req.title, "body": "",
        "created_at": now, "updated_at": now,
    }


@router.get("/docs/{doc_id}")
async def get_doc(
    doc_id: str,
    current_user: str = Depends(get_current_user),
) -> dict[str, Any]:
    with get_db_connection() as conn:
        conn.row_factory = sqlite3.Row
        _ensure_schema(conn)
        row = _get_doc_or_404(conn, current_user, doc_id)
        refs = conn.execute(
            "SELECT id, ref_type, target_id, granularity, snapshot, "
            "       source_nodes, created_at "
            "FROM doc_references WHERE user_id = ? AND doc_id = ? "
            "ORDER BY created_at ASC",
            (current_user, doc_id),
        ).fetchall()
        return {
            "id": row["id"], "title": row["title"], "body": row["body"],
            "created_at": row["created_at"],
            "updated_at": row["updated_at"],
            "references": [
                {
                    "ref_id":      r["id"],
                    "ref_type":    r["ref_type"],
                    "target_id":   r["target_id"],
                    "granularity": r["granularity"],
                    "chip_label":  _chip_label_for_row(conn, current_user, r),
                    "snapshot":    r["snapshot"],
                    "source_nodes": json.loads(r["source_nodes"] or "[]"),
                    "created_at":  r["created_at"],
                }
                for r in refs
            ],
        }


@router.put("/docs/{doc_id}")
async def update_doc(
    doc_id: str,
    req: DocUpdateRequest,
    current_user: str = Depends(get_current_user),
) -> dict[str, Any]:
    with get_db_connection() as conn:
        conn.row_factory = sqlite3.Row
        _ensure_schema(conn)
        row = _get_doc_or_404(conn, current_user, doc_id)
        now = _now()
        new_title = req.title if req.title is not None else row["title"]
        body_changed = req.body is not None and req.body != row["body"]
        new_body = req.body if req.body is not None else row["body"]
        conn.execute(
            "UPDATE docs SET title = ?, body = ?, updated_at = ? "
            "WHERE user_id = ? AND id = ?",
            (new_title, new_body, now, current_user, doc_id),
        )
        snapshot_id: Optional[int] = None
        if body_changed:
            cur = conn.execute(
                "INSERT INTO doc_snapshots "
                "(doc_id, user_id, body, label, created_at) "
                "VALUES (?, ?, ?, ?, ?)",
                (doc_id, current_user, new_body, req.snapshot_label, now),
            )
            snapshot_id = cur.lastrowid
            # Retention cap — keep the latest _SNAPSHOT_CAP per doc.
            conn.execute(
                "DELETE FROM doc_snapshots "
                "WHERE user_id = ? AND doc_id = ? AND id NOT IN ("
                "  SELECT id FROM doc_snapshots "
                "  WHERE user_id = ? AND doc_id = ? "
                "  ORDER BY id DESC LIMIT ?)",
                (current_user, doc_id, current_user, doc_id, _SNAPSHOT_CAP),
            )
        conn.commit()
        return {
            "ok": True, "updated_at": now,
            "snapshot_created": body_changed,
            "snapshot_id": snapshot_id,
        }


@router.delete("/docs/{doc_id}")
async def delete_doc(
    doc_id: str,
    current_user: str = Depends(get_current_user),
) -> dict[str, Any]:
    with get_db_connection() as conn:
        conn.row_factory = sqlite3.Row
        _ensure_schema(conn)
        _get_doc_or_404(conn, current_user, doc_id)
        conn.execute(
            "DELETE FROM doc_references WHERE user_id = ? AND doc_id = ?",
            (current_user, doc_id),
        )
        conn.execute(
            "DELETE FROM doc_snapshots WHERE user_id = ? AND doc_id = ?",
            (current_user, doc_id),
        )
        conn.execute(
            "DELETE FROM doc_chat_messages WHERE user_id = ? AND doc_id = ?",
            (current_user, doc_id),
        )
        conn.execute(
            "DELETE FROM docs WHERE user_id = ? AND id = ?",
            (current_user, doc_id),
        )
        conn.commit()
    return {"ok": True}


# ─────────────────────────────────────────────────────────────────────
# Version snapshots
# ─────────────────────────────────────────────────────────────────────

@router.get("/docs/{doc_id}/snapshots")
async def list_snapshots(
    doc_id: str,
    current_user: str = Depends(get_current_user),
) -> dict[str, Any]:
    with get_db_connection() as conn:
        conn.row_factory = sqlite3.Row
        _ensure_schema(conn)
        _get_doc_or_404(conn, current_user, doc_id)
        rows = conn.execute(
            "SELECT id, label, created_at, LENGTH(body) AS chars "
            "FROM doc_snapshots WHERE user_id = ? AND doc_id = ? "
            "ORDER BY id DESC",
            (current_user, doc_id),
        ).fetchall()
        return {
            "snapshots": [
                {
                    "id": r["id"], "label": r["label"],
                    "created_at": r["created_at"],
                    "chars": int(r["chars"] or 0),
                }
                for r in rows
            ]
        }


@router.post("/docs/{doc_id}/snapshots/{snapshot_id}/restore")
async def restore_snapshot(
    doc_id: str,
    snapshot_id: int,
    current_user: str = Depends(get_current_user),
) -> dict[str, Any]:
    with get_db_connection() as conn:
        conn.row_factory = sqlite3.Row
        _ensure_schema(conn)
        _get_doc_or_404(conn, current_user, doc_id)
        snap = conn.execute(
            "SELECT id, body FROM doc_snapshots "
            "WHERE user_id = ? AND doc_id = ? AND id = ?",
            (current_user, doc_id, snapshot_id),
        ).fetchone()
        if snap is None:
            raise HTTPException(status_code=404, detail="snapshot not found")
        now = _now()
        conn.execute(
            "UPDATE docs SET body = ?, updated_at = ? "
            "WHERE user_id = ? AND id = ?",
            (snap["body"], now, current_user, doc_id),
        )
        # The restore is itself a version event so the medic can undo
        # the undo.
        conn.execute(
            "INSERT INTO doc_snapshots "
            "(doc_id, user_id, body, label, created_at) "
            "VALUES (?, ?, ?, ?, ?)",
            (doc_id, current_user, snap["body"],
             f"restore:{snapshot_id}", now),
        )
        conn.commit()
        return {"ok": True, "body": snap["body"], "updated_at": now}


# ─────────────────────────────────────────────────────────────────────
# References (@ chips)
# ─────────────────────────────────────────────────────────────────────

@router.post("/docs/{doc_id}/references")
async def create_reference(
    doc_id: str,
    req: ReferenceCreateRequest,
    current_user: str = Depends(get_current_user),
) -> dict[str, Any]:
    """Resolve an @ selection into a FROZEN de-identified snapshot +
    chip metadata, and write the audit event (who/when/doc/target/
    granularity)."""
    with get_db_connection() as conn:
        conn.row_factory = sqlite3.Row
        _ensure_schema(conn)
        _get_doc_or_404(conn, current_user, doc_id)

        snapshot, source_nodes, chip_label = _build_reference_snapshot(
            conn, current_user, req.ref_type, req.target_id,
            req.granularity,
        )

        ref_id = str(uuid.uuid4())
        now = _now()
        conn.execute(
            "INSERT INTO doc_references "
            "(id, doc_id, user_id, ref_type, target_id, granularity, "
            " snapshot, source_nodes, created_at) "
            "VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)",
            (ref_id, doc_id, current_user, req.ref_type, req.target_id,
             req.granularity, snapshot,
             json.dumps(source_nodes, ensure_ascii=False), now),
        )
        conn.commit()

        # Audit event — best-effort but loud on failure. The reference
        # row is the projection; the event is the medico-legal record.
        try:
            from nexus_server.event_sourcing import (
                EventKind, Store, init_event_sourcing_schema,
            )
            init_event_sourcing_schema(conn)
            store = Store(conn)
            store.emit_and_apply(
                kind=EventKind.DOC_REFERENCE_CREATED,
                payload={
                    "doc_id":            doc_id,
                    "ref_id":            ref_id,
                    "ref_type":          req.ref_type,
                    "target_id":         req.target_id,
                    "granularity":       req.granularity,
                    "source_node_count": len(source_nodes),
                },
                apply_fn=lambda *_a, **_k: None,
                user_id=current_user,
                patient_hash=(
                    req.target_id if req.ref_type == "patient" else None
                ),
            )
        except Exception as exc:  # noqa: BLE001
            logger.warning("doc_reference audit emit failed: %s", exc)

        return {
            "ref_id": ref_id,
            "chip_label": chip_label,
            "snapshot_preview": snapshot[:400],
        }


# ─────────────────────────────────────────────────────────────────────
# Polish (SSE)
# ─────────────────────────────────────────────────────────────────────

_POLISH_PERSONA = (
    "你是一名严谨的临床写作编辑，服务于医生用户的病例报告、研究摘要与"
    "伦理材料写作。保持医学术语准确、语气专业；只输出改写后的文本本身，"
    "不要输出解释、前言或 Markdown 代码块。"
)
_POLISH_GROUNDING_RULE = "仅基于提供的引用数据改写，不得编造数值。"


def _extract_numbers(text: str) -> set[str]:
    """Normalised numeric tokens (percent sign stripped) for the
    provenance check."""
    return {m.group(0).rstrip("%") for m in _NUM_RE.finditer(text or "")}


@router.post("/docs/{doc_id}/polish")
async def polish_selection(
    doc_id: str,
    req: PolishRequest,
    current_user: str = Depends(get_current_user),
):
    """Rewrite the selected text per instruction, grounded in the doc's
    reference snapshots. Streams SSE frames:

        {type:'revised_chunk', text}          — revised text, chunked
        {type:'provenance_warning', numbers}  — numbers in the revision
                                                with no source in the
                                                selection or refs
        {type:'done', revised}                — full revised text
        {type:'error', message}               — on failure
    """
    # Resolve context BEFORE starting the stream so a bad doc_id is a
    # clean 404, not an SSE error frame.
    with get_db_connection() as conn:
        conn.row_factory = sqlite3.Row
        _ensure_schema(conn)
        _get_doc_or_404(conn, current_user, doc_id)
        snapshots: list[str] = []
        if req.ref_ids:
            qmarks = ",".join("?" * len(req.ref_ids))
            rows = conn.execute(
                f"SELECT snapshot FROM doc_references "
                f"WHERE user_id = ? AND doc_id = ? AND id IN ({qmarks})",
                (current_user, doc_id, *req.ref_ids),
            ).fetchall()
            snapshots = [str(r["snapshot"] or "") for r in rows]

    system_prompt = _POLISH_PERSONA + "\n" + _POLISH_GROUNDING_RULE
    if req.instruction.strip():
        system_prompt += f"\n改写要求：{req.instruction.strip()}"
    if snapshots:
        system_prompt += (
            "\n\n以下是本文档的引用数据（已脱敏），改写时只能使用这些"
            "数据中的数值：\n"
            + "\n---\n".join(snapshots)
        )

    async def event_stream() -> AsyncIterator[str]:
        try:
            # Late import + module-attribute call so tests can
            # monkeypatch llm_gateway.call_llm (same pattern the chat
            # tests rely on).
            from nexus_server import llm_gateway
            content, model_used, _stop, _tools = await llm_gateway.call_llm(
                [{"role": "user", "content": req.selection}],
                system_prompt,
                None,      # model → DEFAULT_LLM_MODEL via provider dispatch
                0.3,
                4096,
            )
            revised = (content or "").strip()
            if not revised:
                yield _sse({
                    "type": "error",
                    "message": "LLM returned empty revision",
                })
                return

            # Stream the revision in chunks so the frontend can render
            # the diff progressively.
            chunk_size = 120
            for i in range(0, len(revised), chunk_size):
                yield _sse({
                    "type": "revised_chunk",
                    "text": revised[i:i + chunk_size],
                })

            # Hallucinated-value guard: numbers in the revision that
            # appear in neither the selection nor any ref snapshot.
            allowed = _extract_numbers(req.selection)
            for s in snapshots:
                allowed |= _extract_numbers(s)
            suspicious: list[str] = []
            for m in _NUM_RE.finditer(revised):
                token = m.group(0)
                if token.rstrip("%") not in allowed and token not in suspicious:
                    suspicious.append(token)
            if suspicious:
                yield _sse({
                    "type": "provenance_warning",
                    "numbers": suspicious,
                })

            yield _sse({
                "type": "done",
                "revised": revised,
                "model": model_used,
            })
        except Exception as exc:  # noqa: BLE001
            logger.warning("polish stream failed: %s", exc)
            yield _sse({"type": "error", "message": str(exc)[:500]})

    return StreamingResponse(event_stream(), media_type="text/event-stream")


# ─────────────────────────────────────────────────────────────────────
# Conversational co-writing (chat, SSE)
# ─────────────────────────────────────────────────────────────────────
#
# Product pivot: the primary writing interaction is a chat with the AI
# which GENERATES and REVISES the document; the human reviews and
# directs. The model answers conversationally and — only when the turn
# requires a document change — additionally emits the complete updated
# body wrapped in <doc>...</doc>. The server splits the stream into
# reply frames vs doc frames, applies the doc (snapshot-first), and
# persists the transcript in doc_chat_messages.

_CHAT_HISTORY_WINDOW = 12
_CHAT_LLM_CHUNK = 64

_DOC_OPEN = "<doc>"
_DOC_CLOSE = "</doc>"

_CHAT_PERSONA = (
    "你是一名严谨的临床写作副驾驶，与医生用户通过对话协作撰写病例报告、"
    "研究摘要与伦理材料。文档正文由你生成和修改；用户负责审阅并提出要求。"
    "保持医学术语准确、语气专业。"
)

_CHAT_OUTPUT_CONTRACT = (
    "输出规则（必须严格遵守）：\n"
    "1. 始终先输出一段简短的对话回复，面向作者，说明修改思路或回答问题。\n"
    "2. 仅当用户的请求需要修改文档正文时，才在对话回复之后额外输出更新后的"
    "文档正文，并且必须用 <doc> 和 </doc> 标签精确包裹。<doc> 块内必须是"
    "完整的全文正文，绝不能只输出片段、差异或省略号。\n"
    "3. 正文中形如 {{ref:ID}} 的引用占位符必须原样逐字保留，"
    "不得改写、编造或删除。\n"
    "4. 如果用户只是提问、讨论或请求不涉及正文改动，只输出对话回复，"
    "不要输出 <doc> 块。\n"
    "5. 仅基于当前正文与提供的引用数据写作，不得编造数值。"
)


def _partial_tag_suffix_len(buf: str, tag: str) -> int:
    """Length of the longest strict-prefix of ``tag`` that ``buf`` ends
    with — i.e. how many trailing chars might be the start of a tag
    split across stream chunks and must be held back."""
    for k in range(min(len(buf), len(tag) - 1), 0, -1):
        if buf.endswith(tag[:k]):
            return k
    return 0


class _DocTagParser:
    """Incremental splitter of an LLM stream into conversational reply
    text vs <doc>...</doc> body text.

    ``feed`` returns a list of (kind, text) events with kind in
    {'reply', 'doc_started', 'doc'}. Tags split across chunk boundaries
    are handled by holding back any buffer suffix that could still turn
    out to be the start of the tag we're looking for."""

    def __init__(self) -> None:
        self._buf = ""
        self.in_doc = False
        self.doc_seen = False

    def feed(self, chunk: str) -> list[tuple[str, str]]:
        self._buf += chunk
        out: list[tuple[str, str]] = []
        while True:
            tag = _DOC_CLOSE if self.in_doc else _DOC_OPEN
            idx = self._buf.find(tag)
            if idx != -1:
                if idx > 0:
                    out.append(
                        ("doc" if self.in_doc else "reply", self._buf[:idx])
                    )
                self._buf = self._buf[idx + len(tag):]
                if self.in_doc:
                    self.in_doc = False
                else:
                    self.in_doc = True
                    self.doc_seen = True
                    out.append(("doc_started", ""))
                continue
            hold = _partial_tag_suffix_len(self._buf, tag)
            emit_len = len(self._buf) - hold
            if emit_len > 0:
                out.append(
                    ("doc" if self.in_doc else "reply",
                     self._buf[:emit_len])
                )
                self._buf = self._buf[emit_len:]
            return out

    def finish(self) -> list[tuple[str, str]]:
        """Flush whatever is left (held-back partial tag / unterminated
        doc block — an unterminated <doc> still counts as doc text)."""
        if not self._buf:
            return []
        out = [("doc" if self.in_doc else "reply", self._buf)]
        self._buf = ""
        return out


def _strip_unknown_ref_tokens(
    body: str, valid_ids: set[str],
) -> tuple[str, list[str]]:
    """Defensive guard: drop any {{ref:ID}} token whose ID is not one of
    this doc's references (the model must never invent chips). Returns
    (cleaned_body, removed_ids)."""
    removed: list[str] = []

    def _sub(m: re.Match) -> str:
        rid = m.group(1)
        if rid in valid_ids or rid.lower() in valid_ids:
            return m.group(0)
        removed.append(rid)
        return ""

    return _REF_PLACEHOLDER_RE.sub(_sub, body), removed


@router.get("/docs/{doc_id}/chat")
async def get_doc_chat(
    doc_id: str,
    current_user: str = Depends(get_current_user),
) -> dict[str, Any]:
    """Chronological co-writing transcript for this doc."""
    with get_db_connection() as conn:
        conn.row_factory = sqlite3.Row
        _ensure_schema(conn)
        _get_doc_or_404(conn, current_user, doc_id)
        rows = conn.execute(
            "SELECT id, role, text, doc_applied, created_at "
            "FROM doc_chat_messages WHERE user_id = ? AND doc_id = ? "
            "ORDER BY created_at ASC, rowid ASC",
            (current_user, doc_id),
        ).fetchall()
        return {
            "messages": [
                {
                    "id": r["id"], "role": r["role"], "text": r["text"],
                    "doc_applied": int(r["doc_applied"] or 0),
                    "created_at": r["created_at"],
                }
                for r in rows
            ]
        }


@router.post("/docs/{doc_id}/chat")
async def doc_chat(
    doc_id: str,
    req: ChatRequest,
    current_user: str = Depends(get_current_user),
):
    """One co-writing turn. Streams SSE frames:

        {type:'reply_chunk', text}            — conversational reply
        {type:'doc_started'}                  — <doc> block opened
        {type:'doc_chunk', text}              — doc body text
        {type:'provenance_warning', numbers}  — numbers in the new body
                                                with no source in the
                                                previous body / refs /
                                                user message
        {type:'done', reply, doc_body|null,
                      snapshot_id|null}       — turn complete
        {type:'error', message}               — LLM failure (the user
                                                message is still saved)

    When the model emits a <doc> block the PREVIOUS body is snapshotted
    first (label '对话修订前'), then the doc is updated PUT-style and the
    assistant message is stored with doc_applied=1.
    """
    # Resolve all context BEFORE starting the stream so a bad doc_id is
    # a clean 404, not an SSE error frame.
    with get_db_connection() as conn:
        conn.row_factory = sqlite3.Row
        _ensure_schema(conn)
        row = _get_doc_or_404(conn, current_user, doc_id)
        title = str(row["title"] or "")
        prev_body = str(row["body"] or "")

        ref_rows = conn.execute(
            "SELECT id, snapshot FROM doc_references "
            "WHERE user_id = ? AND doc_id = ? ORDER BY created_at ASC",
            (current_user, doc_id),
        ).fetchall()
        valid_ref_ids = {str(r["id"]) for r in ref_rows}
        if req.ref_ids:
            wanted = set(req.ref_ids)
            snapshots = [
                str(r["snapshot"] or "") for r in ref_rows
                if str(r["id"]) in wanted
            ]
        else:
            snapshots = [str(r["snapshot"] or "") for r in ref_rows]

        # ── ATTACHMENTS ──────────────────────────────────────────────
        # Same resolution as chat_router (the v2 chat path): user-scoped
        # ``uploads`` lookup by file_id, cached ``extracted_text`` with
        # an on-demand extraction fallback (nexus_core distiller via
        # files._bytes_to_text, cached back to the row), unknown /
        # foreign file_ids silently skipped. Each attachment's inlined
        # text is capped at the same 8 KB the v2 path uses so a
        # 500-page PDF can't blow the prompt context.
        attachment_sections: list[tuple[str, str]] = []
        for att in req.attachments or []:
            if not isinstance(att, dict):
                continue
            fid = str(att.get("file_id") or "").strip()
            if not fid:
                continue
            try:
                arow = conn.execute(
                    "SELECT name, mime, extracted_text, disk_path "
                    "FROM uploads "
                    "WHERE user_id = ? AND file_id = ?",
                    (current_user, fid),
                ).fetchone()
            except Exception:  # noqa: BLE001
                arow = None
            if not arow:
                continue
            a_name = str(arow["name"] or att.get("name") or fid)
            a_mime = str(arow["mime"] or "")
            a_text = str(arow["extracted_text"] or "").strip()
            a_path = str(arow["disk_path"] or "")
            a_is_image = a_mime.startswith("image/")

            # On-demand text extraction if not cached (mirrors the
            # Track-A lazy extract in chat_router).
            if not a_text and not a_is_image and a_path:
                try:
                    from pathlib import Path as _Path
                    p = _Path(a_path)
                    if p.is_file():
                        raw = p.read_bytes()
                        from nexus_server.files import (
                            _bytes_to_text, _save_extracted_text,
                        )
                        text_out = _bytes_to_text(raw, a_name, a_mime)
                        if text_out:
                            a_text = text_out.strip()
                            _save_extracted_text(fid, a_text)
                except Exception as e:  # noqa: BLE001
                    logger.debug(
                        "doc chat: lazy extract for %s failed: %s",
                        fid[:8], e,
                    )

            if a_text:
                section = a_text[:8000]
            elif a_is_image:
                section = (
                    "（图片附件——未提取到文字内容。写作对话暂不支持直接"
                    "查看图片；如写作依赖其内容，请向用户说明。）"
                )
            else:
                section = (
                    "（二进制文件——无法提取文本内容。如写作依赖其内容，"
                    "请向用户说明并请求文字版。）"
                )
            attachment_sections.append((a_name, section))

        # Last N turns, chronological, EXCLUDING the message being sent
        # now (it's appended below as the live user turn).
        hist_rows = conn.execute(
            "SELECT role, text FROM doc_chat_messages "
            "WHERE user_id = ? AND doc_id = ? "
            "ORDER BY created_at DESC, rowid DESC LIMIT ?",
            (current_user, doc_id, _CHAT_HISTORY_WINDOW),
        ).fetchall()
        history = [
            {"role": str(r["role"]), "content": str(r["text"] or "")}
            for r in reversed(hist_rows)
        ]

        # Persist the user message immediately — it must survive even
        # if the LLM call fails.
        user_msg_id = str(uuid.uuid4())
        conn.execute(
            "INSERT INTO doc_chat_messages "
            "(id, doc_id, user_id, role, text, doc_applied, created_at) "
            "VALUES (?, ?, ?, 'user', ?, 0, ?)",
            (user_msg_id, doc_id, current_user, req.message, _now()),
        )
        conn.commit()

    system_prompt = (
        _CHAT_PERSONA + "\n\n" + _CHAT_OUTPUT_CONTRACT
        + f"\n\n文档标题：{title or '（未命名）'}"
        + "\n当前文档正文（{{ref:ID}} 为引用占位符，必须原样保留）：\n"
        + (prev_body if prev_body else "（正文为空）")
    )
    if snapshots:
        system_prompt += (
            "\n\n以下是本文档的引用数据（已脱敏），写作时只能使用这些"
            "数据中的数值：\n" + "\n---\n".join(snapshots)
        )

    # ── USER ATTACHMENTS ─────────────────────────────────────────────
    # Per-turn 📎 uploads from the writing composer. Injected AFTER the
    # reference snapshots (they are additional source material, ranked
    # below the doc's curated references) and BEFORE the skills block
    # (skills may override tone/format and must stay last).
    for a_name, a_text in attachment_sections:
        system_prompt += f"\n\n## 用户附件: {a_name}\n{a_text}"

    # ── ACTIVE SKILLS ────────────────────────────────────────────────
    # Same injection as chat_router: explicit "/" invocations from
    # req.skills (installed+enabled only; others silently dropped) +
    # every enabled auto_apply skill. Appended LAST so skill
    # instructions can override tone/format defaults without touching
    # the co-writing output contract above. Non-fatal on failure.
    try:
        from nexus_server.skills_router import build_skills_block
        skills_block, _applied_skills = build_skills_block(
            current_user, req.skills or [],
        )
        if skills_block:
            system_prompt += "\n\n" + skills_block
    except Exception as exc:  # noqa: BLE001
        logger.warning(
            "doc chat: skills block build failed (non-fatal): %s", exc,
        )

    messages = history + [{"role": "user", "content": req.message}]

    async def event_stream() -> AsyncIterator[str]:
        try:
            # Same dispatch helper as polish — late import + module
            # attribute call so tests can monkeypatch
            # llm_gateway.call_llm; provider routing (kimi/openai/
            # anthropic/gemini) happens inside call_llm.
            from nexus_server import llm_gateway
            content, model_used, _stop, _tools = await llm_gateway.call_llm(
                messages,
                system_prompt,
                None,      # model → DEFAULT_LLM_MODEL via provider dispatch
                0.3,
                8192,
            )
            content = (content or "").strip()
            if not content:
                yield _sse({
                    "type": "error",
                    "message": "LLM returned empty response",
                })
                return

            # Incremental parse: feed the model output through the tag
            # splitter in chunks (tags may straddle chunk boundaries).
            parser = _DocTagParser()
            reply_parts: list[str] = []
            doc_parts: list[str] = []

            def _frames(events: list[tuple[str, str]]) -> list[str]:
                out: list[str] = []
                for kind, text in events:
                    if kind == "reply":
                        reply_parts.append(text)
                        out.append(_sse({
                            "type": "reply_chunk", "text": text,
                        }))
                    elif kind == "doc_started":
                        out.append(_sse({"type": "doc_started"}))
                    else:  # 'doc'
                        doc_parts.append(text)
                        out.append(_sse({
                            "type": "doc_chunk", "text": text,
                        }))
                return out

            for i in range(0, len(content), _CHAT_LLM_CHUNK):
                for frame in _frames(
                    parser.feed(content[i:i + _CHAT_LLM_CHUNK])
                ):
                    yield frame
            for frame in _frames(parser.finish()):
                yield frame

            reply = "".join(reply_parts).strip()
            doc_body: Optional[str] = (
                "".join(doc_parts).strip() if parser.doc_seen else None
            )

            snapshot_id: Optional[int] = None
            now = _now()
            assistant_msg_id = str(uuid.uuid4())

            if doc_body is not None:
                # Defensive: the model must not invent reference chips.
                doc_body, removed = _strip_unknown_ref_tokens(
                    doc_body, valid_ref_ids,
                )
                if removed:
                    logger.warning(
                        "doc chat: stripped %d unknown ref token(s): %s",
                        len(removed), removed,
                    )

                # Hallucinated-value guard: numbers in the new body
                # that appear nowhere in the previous body, the ref
                # snapshots injected as context, or the user message.
                allowed = _extract_numbers(prev_body)
                allowed |= _extract_numbers(req.message)
                for s in snapshots:
                    allowed |= _extract_numbers(s)
                # Attachment text is legitimate source material the
                # medic supplied this turn — numbers copied from it
                # must not trip the hallucination guard.
                for _a_name, a_text in attachment_sections:
                    allowed |= _extract_numbers(a_text)
                suspicious: list[str] = []
                for m in _NUM_RE.finditer(doc_body):
                    token = m.group(0)
                    if (token.rstrip("%") not in allowed
                            and token not in suspicious):
                        suspicious.append(token)
                if suspicious:
                    yield _sse({
                        "type": "provenance_warning",
                        "numbers": suspicious,
                    })

                with get_db_connection() as conn:
                    conn.row_factory = sqlite3.Row
                    _ensure_schema(conn)
                    # Snapshot the PREVIOUS body first so the author
                    # can always revert the AI revision.
                    cur = conn.execute(
                        "INSERT INTO doc_snapshots "
                        "(doc_id, user_id, body, label, created_at) "
                        "VALUES (?, ?, ?, ?, ?)",
                        (doc_id, current_user, prev_body,
                         "对话修订前", now),
                    )
                    snapshot_id = cur.lastrowid
                    conn.execute(
                        "DELETE FROM doc_snapshots "
                        "WHERE user_id = ? AND doc_id = ? AND id NOT IN ("
                        "  SELECT id FROM doc_snapshots "
                        "  WHERE user_id = ? AND doc_id = ? "
                        "  ORDER BY id DESC LIMIT ?)",
                        (current_user, doc_id, current_user, doc_id,
                         _SNAPSHOT_CAP),
                    )
                    # PUT-equivalent update of the doc body.
                    conn.execute(
                        "UPDATE docs SET body = ?, updated_at = ? "
                        "WHERE user_id = ? AND id = ?",
                        (doc_body, now, current_user, doc_id),
                    )
                    conn.execute(
                        "INSERT INTO doc_chat_messages "
                        "(id, doc_id, user_id, role, text, doc_applied, "
                        " created_at) "
                        "VALUES (?, ?, ?, 'assistant', ?, 1, ?)",
                        (assistant_msg_id, doc_id, current_user, reply,
                         now),
                    )
                    conn.commit()
            else:
                with get_db_connection() as conn:
                    _ensure_schema(conn)
                    conn.execute(
                        "INSERT INTO doc_chat_messages "
                        "(id, doc_id, user_id, role, text, doc_applied, "
                        " created_at) "
                        "VALUES (?, ?, ?, 'assistant', ?, 0, ?)",
                        (assistant_msg_id, doc_id, current_user, reply,
                         now),
                    )
                    conn.commit()

            yield _sse({
                "type": "done",
                "reply": reply,
                "doc_body": doc_body,
                "snapshot_id": snapshot_id,
                "message_id": assistant_msg_id,
                "model": model_used,
            })
        except Exception as exc:  # noqa: BLE001
            logger.warning("doc chat stream failed: %s", exc)
            yield _sse({"type": "error", "message": str(exc)[:500]})

    return StreamingResponse(event_stream(), media_type="text/event-stream")


# ─────────────────────────────────────────────────────────────────────
# PHI scan
# ─────────────────────────────────────────────────────────────────────

@router.post("/docs/{doc_id}/phi-scan")
async def phi_scan(
    doc_id: str,
    req: Optional[PhiScanRequest] = None,
    current_user: str = Depends(get_current_user),
) -> dict[str, Any]:
    with get_db_connection() as conn:
        conn.row_factory = sqlite3.Row
        _ensure_schema(conn)
        row = _get_doc_or_404(conn, current_user, doc_id)
        override = req.body if req is not None else None
        body = override if override is not None else str(row["body"] or "")
        return {"findings": _phi_scan_body(conn, current_user, body)}


# ─────────────────────────────────────────────────────────────────────
# Export
# ─────────────────────────────────────────────────────────────────────

def _resolution_matches(finding: dict, resolution: dict) -> bool:
    if resolution.get("kind") and resolution["kind"] != finding["kind"]:
        return False
    if (
        resolution.get("start") is not None
        and resolution.get("end") is not None
    ):
        return (
            int(resolution["start"]) == finding["start"]
            and int(resolution["end"]) == finding["end"]
        )
    excerpt = resolution.get("excerpt")
    return bool(excerpt) and excerpt == finding["excerpt"]


@router.post("/docs/{doc_id}/export")
async def export_doc(
    doc_id: str,
    req: ExportRequest,
    current_user: str = Depends(get_current_user),
):
    """PHI gate → chip expansion → .docx.

    422 {code:'phi_unresolved', findings:[...]} when the body still has
    unresolved PHI findings. Resolutions with a ``replacement`` string
    rewrite the flagged span before export; without one they count as
    an explicit keep.
    """
    with get_db_connection() as conn:
        conn.row_factory = sqlite3.Row
        _ensure_schema(conn)
        row = _get_doc_or_404(conn, current_user, doc_id)
        title = str(row["title"] or "")
        body = str(row["body"] or "")

        findings = _phi_scan_body(conn, current_user, body)
        unresolved: list[dict] = []
        resolved: list[tuple[dict, dict]] = []
        for f in findings:
            match = next(
                (r for r in req.resolutions if _resolution_matches(f, r)),
                None,
            )
            if match is None:
                unresolved.append(f)
            else:
                resolved.append((f, match))
        if unresolved:
            return JSONResponse(
                status_code=422,
                content={"code": "phi_unresolved", "findings": unresolved},
            )

        # Apply replacements back-to-front so earlier spans stay valid.
        for f, r in sorted(resolved, key=lambda p: p[0]["start"],
                           reverse=True):
            replacement = r.get("replacement")
            if isinstance(replacement, str):
                body = body[:f["start"]] + replacement + body[f["end"]:]

        # Expand {{ref:ID}} chips to their frozen de-identified text.
        refs = conn.execute(
            "SELECT id, ref_type, target_id, granularity, snapshot, "
            "       source_nodes, created_at "
            "FROM doc_references WHERE user_id = ? AND doc_id = ? "
            "ORDER BY created_at ASC",
            (current_user, doc_id),
        ).fetchall()
        ref_by_id = {str(r["id"]): r for r in refs}

        def _expand(m: re.Match) -> str:
            r = ref_by_id.get(m.group(1))
            return str(r["snapshot"]) if r is not None else ""

        expanded = _REF_PLACEHOLDER_RE.sub(_expand, body)

        # Build the .docx.
        from docx import Document
        document = Document()
        document.add_heading(title or "未命名文档", level=1)
        for para in re.split(r"\n\s*\n", expanded):
            para = para.strip()
            if para:
                document.add_paragraph(para)

        if req.include_sources and refs:
            document.add_heading("引用来源", level=2)
            for r in refs:
                label = _chip_label_for_row(conn, current_user, r)
                try:
                    nodes = json.loads(r["source_nodes"] or "[]")
                except (TypeError, ValueError):
                    nodes = []
                node_str = (
                    ", ".join(str(n) for n in nodes) if nodes else "—"
                )
                document.add_paragraph(
                    f"{label}（{r['ref_type']}/{r['granularity']}）"
                    f" · 来源节点: {node_str}",
                    style="List Bullet",
                )

        buf = BytesIO()
        document.save(buf)

    filename = f"doc-{doc_id[:8]}.docx"
    return Response(
        content=buf.getvalue(),
        media_type=(
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.document"
        ),
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
        },
    )
