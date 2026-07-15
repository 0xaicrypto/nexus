"""Generic file distillation: text extraction + LLM-driven summarization.

Originally lived in ``nexus_server.attachment_distiller`` — moved here so
any consumer of the SDK can use the same pipeline without re-implementing
it. The server-side persistence (writing the resulting summary into
``sync_events`` so it rides the event log + BSC anchors) is intentionally
NOT in this module: that's a deployment-specific concern that the
server keeps in its own thin shim. This module only does:

    raw bytes / base64 / text  →  extract_text()  →  distill()  →  summary

Persistence, request shape, and storage are all caller-owned.

Why split this way:
  * Server hosts a multi-tenant chat where attachments are folded into
    the prompt + persisted into per-user audit logs.
  * Future P2P / standalone twin agents will want the same "summarise
    this file before stuffing it into the prompt" behaviour without any
    of the server's HTTP / SQL infrastructure.
  * Tests don't need to spin up a database to verify summary quality.

Soft dependencies:
  * ``pypdf`` for PDF text extraction. If it's missing, PDF inputs fall
    back to a metadata stub — the LLM still gets a useful description
    via filename/mime, just no body text.
"""

from __future__ import annotations

import base64
import binascii
import io
import logging
from typing import Any, Awaitable, Callable, Optional

logger = logging.getLogger(__name__)


# ── Tunables ──────────────────────────────────────────────────────────

# Cap on how many characters of file text we send to the distiller LLM.
# Large enough to cover most papers + small books; the rest the model
# summarises from the head it sees. Keeps token cost predictable.
DISTILL_INPUT_CHAR_BUDGET = 60_000

# Cap on the distilled summary itself — what shows up in subsequent
# prompts. Roughly 1k tokens.
DISTILL_OUTPUT_CHAR_BUDGET = 4_000


DISTILL_SYSTEM_PROMPT = (
    "You are a careful file summarizer. Given a single file's contents, "
    "produce a structured summary the original user can rely on later.\n\n"
    "Output format (markdown, plain text — NO code fences):\n"
    "- One-line description of what this file is.\n"
    "- 'Key points:' bulleted list (5–12 items): the most important "
    "facts, claims, decisions, or data the reader would want preserved.\n"
    "- 'Entities:' inline list of important named entities (people, "
    "products, dates, places, IDs) — comma separated.\n"
    "- 'Structure:' brief note on how the document is organized.\n"
    "- If the file is mostly tabular, list the column names and one or "
    "two representative rows.\n"
    "- Be concrete — quote short fragments where helpful.\n"
    f"- Hard limit: {DISTILL_OUTPUT_CHAR_BUDGET} characters total.\n"
    "- If you cannot read the content (binary or empty), still describe "
    "what you can infer from the filename and mime type."
)


# #128 — image caption distill. Separate prompt because vision input
# is fundamentally different from text: we want structured fields
# the memory layer can index, not free-form prose. The output is
# intentionally short (~600 chars) so we don't bloat attachment_distilled
# rows that will be re-read every session via the uploads memory block.
IMAGE_DISTILL_SYSTEM_PROMPT = (
    "You are a visual indexer for a long-running personal assistant. "
    "Given a single image, produce a STRUCTURED CAPTION that the "
    "assistant can use months later to remember what the user showed "
    "it — without re-reading the original pixels.\n\n"
    "Output format (markdown, plain text — NO code fences, NO yaml "
    "wrapping):\n"
    "- 'kind:' one of {screenshot, photo, chart, diagram, code, "
    "  document_scan, ui_mockup, medical_imaging, art, meme, other}.\n"
    "- 'domain:' one short phrase for the subject area (e.g. "
    "  'crypto trading view', 'github pull request', "
    "  'chest CT axial', 'family vacation photo', 'react component', "
    "  'shopping receipt'). Be concrete; this is what the assistant "
    "  will grep on to find this image later.\n"
    "- 'summary:' 1–3 sentences describing what's visually in the "
    "  image. Focus on layout, content, and key text that's actually "
    "  visible — not interpretation.\n"
    "- 'salient_text:' the most important short text strings visible "
    "  in the image (button labels, headers, prices, error messages, "
    "  identifiers, error codes). Comma-separated. Skip if no readable "
    "  text. Do NOT transcribe entire paragraphs — caller can OCR if "
    "  they need that.\n"
    "- 'entities:' people / products / orgs / tickers / file names "
    "  visible. Comma-separated; skip if none.\n"
    "- 'follow_up_hints:' 1–2 short phrases describing what kind of "
    "  questions the user likely wants help with based on this image "
    "  (e.g. 'analyze support/resistance', 'review code style', "
    "  'identify radiological finding'). Helps the assistant pick a "
    "  default response framing.\n"
    "- Be literal. If you can't tell what something is, say "
    "  'unclear' instead of guessing.\n"
    "- Hard limit: 800 characters total. Image captions should be "
    "  scannable in a memory list, not encyclopedic."
)


# ── Type alias for an LLM caller ──────────────────────────────────────
#
# Callers thread their own LLM in. The signature mirrors the server's
# ``llm_gateway.call_llm`` historical contract (which both server and
# twin/SDK callers can adapt to):
#
#   async fn(messages, system_prompt, model, temperature, max_tokens,
#            tools) -> (content, model_used, stop_reason, tool_calls)
#
# The first arg is a list of {"role", "content"} dicts. ``tools`` is
# always None for distillation — this is a leaf call, no tool use.

LlmFn = Callable[
    [
        list[dict],          # messages
        Optional[str],       # system_prompt
        Optional[str],       # model (None → caller's default)
        Optional[float],     # temperature
        Optional[int],       # max_tokens
        Any,                 # tools (always None for distillation)
    ],
    Awaitable[tuple[str, str, str, list]],  # (content, model, stop, tool_calls)
]


# ── Text extraction ───────────────────────────────────────────────────


def extract_text(
    name: str,
    mime: str,
    content_text: Optional[str],
    content_base64: Optional[str],
) -> tuple[str, str]:
    """Pull plain text out of an attachment for the distiller.

    Returns ``(text, source_label)`` where ``source_label`` tells
    downstream code *how* the text was obtained ("text" / "pdf" /
    "binary-stub" / "empty"). Always returns *something* — never raises
    — so callers can assume a distill attempt is always possible.

    Args:
        name: Filename — used in the no-content stub messages.
        mime: MIME type — used to pick the extraction path.
        content_text: Pre-decoded UTF-8 content (preferred when
            available; the caller did the decode).
        content_base64: Raw bytes encoded as base64. Used as fallback.
    """
    if content_text is not None and content_text:
        return content_text[:DISTILL_INPUT_CHAR_BUDGET], "text"

    if content_base64 is None:
        return f"[empty attachment named {name!r} ({mime})]", "empty"

    try:
        raw = base64.b64decode(content_base64, validate=False)
    except (binascii.Error, ValueError) as e:
        return (
            f"[unreadable attachment {name!r}: bad base64 ({e})]",
            "binary-stub",
        )

    # PDF: try pypdf if it's installed; not a hard dependency.
    if mime == "application/pdf" or name.lower().endswith(".pdf"):
        text = _extract_pdf_text(raw)
        if text:
            return text[:DISTILL_INPUT_CHAR_BUDGET], "pdf"
        return (
            f"[PDF {name!r} — text extraction unavailable; "
            f"{len(raw)} bytes]",
            "binary-stub",
        )

    # DOCX: Word documents. ZIP wrappers around word/document.xml.
    # Try python-docx (clean paragraph extraction) first, then fall
    # back to a stdlib zipfile + lxml/regex parser so we still work
    # when the SDK is installed without the optional ``docx`` extra.
    _DOCX_MIME = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",  # legacy .doc — extraction may fail but try
    )
    if mime in _DOCX_MIME or name.lower().endswith((".docx", ".doc")):
        text = _extract_docx_text(raw)
        if text:
            return text[:DISTILL_INPUT_CHAR_BUDGET], "docx"
        return (
            f"[DOCX {name!r} — text extraction unavailable; "
            f"{len(raw)} bytes]",
            "binary-stub",
        )

    # Plain text dressed up as binary by an over-cautious client?
    try:
        decoded = raw.decode("utf-8")
        return decoded[:DISTILL_INPUT_CHAR_BUDGET], "text"
    except UnicodeDecodeError as exc:
        logger.debug("input is not UTF-8 text: %s", exc)

    return (
        f"[binary attachment {name!r} ({mime}, {len(raw)} bytes); "
        f"content not extracted]",
        "binary-stub",
    )


def _extract_pdf_text(raw: bytes) -> str:
    """Try to pull text out of a PDF. Returns empty string if pypdf
    isn't available or extraction fails — the caller falls back to a
    metadata-only summary.
    """
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError:
        logger.debug("pypdf not installed; skipping PDF text extraction")
        return ""

    try:
        reader = PdfReader(io.BytesIO(raw))
        chunks: list[str] = []
        running_total = 0
        for i, page in enumerate(reader.pages):
            try:
                t = page.extract_text() or ""
            except Exception:
                t = ""
            if not t:
                continue
            chunks.append(f"\n--- Page {i + 1} ---\n{t}")
            running_total += len(t)
            if running_total >= DISTILL_INPUT_CHAR_BUDGET:
                break
        return "".join(chunks)
    except Exception as e:
        logger.debug("PDF extraction failed: %s", e)
        return ""


def _extract_docx_text(raw: bytes) -> str:
    """Pull text out of a Word .docx file.

    Strategy (best → fallback):
      1. ``python-docx`` — clean paragraph + table walk. Preferred,
         but optional dependency.
      2. ``zipfile`` + xml.etree — stdlib-only fallback. .docx is just
         a ZIP archive whose ``word/document.xml`` holds the text in
         ``<w:t>`` runs. This works even when python-docx isn't
         installed (eg. on a minimal user machine).

    Headings / tables / lists all flatten to plain paragraphs joined
    by blank lines. That's fine for the distiller — it cares about
    semantic content, not Word formatting.

    Returns "" on total failure; the caller renders a metadata stub.
    """
    # ── Path 1: python-docx ────────────────────────────────────────
    try:
        import docx as _docx  # type: ignore
    except ImportError:
        _docx = None  # type: ignore

    if _docx is not None:
        try:
            doc = _docx.Document(io.BytesIO(raw))
            chunks: list[str] = []
            running_total = 0

            # Paragraphs in document order.
            for p in doc.paragraphs:
                t = (p.text or "").strip()
                if not t:
                    continue
                chunks.append(t)
                running_total += len(t)
                if running_total >= DISTILL_INPUT_CHAR_BUDGET:
                    return "\n\n".join(chunks)

            # Tables — flatten cells row-by-row so column relationships
            # survive in plain text. Useful for spec docs / contracts /
            # data sheets that put real content in tables.
            for tbl in doc.tables:
                for row in tbl.rows:
                    row_text = " | ".join(
                        (c.text or "").strip() for c in row.cells
                    )
                    if row_text.strip(" |"):
                        chunks.append(row_text)
                        running_total += len(row_text)
                        if running_total >= DISTILL_INPUT_CHAR_BUDGET:
                            return "\n\n".join(chunks)

            return "\n\n".join(chunks)
        except Exception as e:
            logger.debug("python-docx extraction failed, falling back: %s", e)

    # ── Path 2: stdlib zipfile + xml.etree ─────────────────────────
    try:
        import xml.etree.ElementTree as ET
        import zipfile

        with zipfile.ZipFile(io.BytesIO(raw)) as zf:
            try:
                xml_bytes = zf.read("word/document.xml")
            except KeyError:
                logger.debug("docx missing word/document.xml — not a valid Word file?")
                return ""

        # WordprocessingML namespace. Hard-code it to avoid xmlns gymnastics.
        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        root = ET.fromstring(xml_bytes)

        chunks: list[str] = []
        running_total = 0
        # Iterate over paragraphs; concatenate every <w:t> child of
        # every <w:r> run inside. Drops formatting but preserves
        # paragraph boundaries — same as Word's "save as text" does.
        for para in root.iter(f"{ns}p"):
            parts: list[str] = []
            for t in para.iter(f"{ns}t"):
                if t.text:
                    parts.append(t.text)
            line = "".join(parts).strip()
            if not line:
                continue
            chunks.append(line)
            running_total += len(line)
            if running_total >= DISTILL_INPUT_CHAR_BUDGET:
                break
        return "\n\n".join(chunks)
    except Exception as e:
        logger.debug("docx stdlib extraction failed: %s", e)
        return ""


# ── Distillation (LLM-driven) ─────────────────────────────────────────


async def distill(
    *,
    name: str,
    mime: str,
    size_bytes: int,
    content_text: Optional[str],
    content_base64: Optional[str],
    llm_fn: LlmFn,
) -> tuple[str, str]:
    """Distill a single attachment via the LLM.

    Returns ``(summary, source_label)``. On any LLM failure, falls
    back to a head-truncation of whatever text we extracted, so the
    caller is never blocked. ``source_label`` includes a ``+fallback``
    suffix when the LLM call failed and the head excerpt was used.
    """
    text, source = extract_text(name, mime, content_text, content_base64)

    user_msg = (
        f"Filename: {name}\nMIME: {mime}\nSize: {size_bytes} bytes\n"
        f"Source: {source}\n\n"
        f"--- begin content ---\n{text}\n--- end content ---"
    )

    try:
        content, _model, _stop, _tools = await llm_fn(
            [{"role": "user", "content": user_msg}],
            DISTILL_SYSTEM_PROMPT,
            None,    # default model
            0.2,     # low temperature: we want consistent factual summaries
            1024,    # output budget
            None,    # no tools — this is a leaf LLM call
        )
        summary = (content or "").strip()
        if summary:
            if len(summary) > DISTILL_OUTPUT_CHAR_BUDGET:
                summary = summary[:DISTILL_OUTPUT_CHAR_BUDGET] + "…"
            return summary, source
    except Exception as e:
        logger.warning("Distillation LLM call failed for %s: %s", name, e)

    # LLM unavailable or returned nothing — fall back to a head excerpt
    # so the caller still has *something* to fold into the chat prompt.
    head = text[:1024]
    fallback = (
        f"[Distillation unavailable; head-only excerpt of {name} "
        f"({mime}, {size_bytes} bytes)]\n\n{head}"
    )
    return fallback, source + "+fallback"


# Cap on the image caption summary — kept tight (vs DISTILL_OUTPUT_CHAR_BUDGET
# at 4k) because image captions sit in the per-session memory block that
# gets re-read every turn. A bloated caption would steal context budget
# from the rest of the conversation. The system prompt enforces ~800
# chars; this is just the hard ceiling.
IMAGE_CAPTION_CHAR_BUDGET = 1200


async def distill_image(
    *,
    name: str,
    mime: str,
    size_bytes: int,
    content_base64: str,
    llm_fn: LlmFn,
) -> tuple[str, str]:
    """Distill a single image attachment via a vision LLM call (#128).

    Mirrors :func:`distill` but routes through the vision multimodal
    path: we attach the raw image bytes as an ``inline_data`` part on
    the user message (via the ``images`` field) and ask the LLM for
    a structured caption with the fields IMAGE_DISTILL_SYSTEM_PROMPT
    specifies (kind / domain / summary / salient_text / entities /
    follow_up_hints).

    Returns ``(caption, source_label)``. On any failure the caller
    gets a neutral stub so the memory chain has *something* to write
    instead of "" — the source label tells downstream code which
    branch fired (``"vision-caption"`` /
    ``"vision-caption+fallback"``).

    Why vision rather than OCR-then-distill: the goal here is
    indexability — caption is read by FTS / vector search later when
    the user asks "the chest CT I showed you last week". A vision
    model sees layout, colour, and content in one pass; OCR misses
    everything except text and loses the modality cue.

    Cost: one extra vision call per uploaded image (≈ ¥0.003 with
    Gemini Flash). Paid once on upload; saved many times over by
    skipping image bytes on subsequent turns referencing the same
    attachment.
    """
    if not content_base64:
        return (
            f"[Image — {name} ({mime}, {size_bytes} bytes); no bytes "
            f"available to caption.]",
            "vision-caption+empty",
        )

    user_msg = {
        "role": "user",
        "content": (
            f"Filename: {name}\n"
            f"MIME: {mime}\n"
            f"Size: {size_bytes} bytes\n\n"
            "Caption this image using the structured fields defined "
            "in your instructions. Stay within the 800-character budget."
        ),
        # #123-shape: the LLM client converts this into a Gemini
        # inline_data Blob part. Other providers (OpenAI vision /
        # Anthropic vision) wire their own image part shape; the
        # llm_fn caller decides which provider is active.
        "images": [{"mime": mime, "data_b64": content_base64}],
    }

    try:
        content, _model, _stop, _tools = await llm_fn(
            [user_msg],
            IMAGE_DISTILL_SYSTEM_PROMPT,
            None,    # default vision-capable model (Gemini 2.5 Flash by default)
            0.2,     # low temp: captions should be repeatable on re-distill
            512,     # output budget — caption is short by design
            None,    # no tools
        )
        caption = (content or "").strip()
        if caption:
            if len(caption) > IMAGE_CAPTION_CHAR_BUDGET:
                caption = caption[:IMAGE_CAPTION_CHAR_BUDGET] + "…"
            return caption, "vision-caption"
    except Exception as e:  # noqa: BLE001
        logger.warning(
            "Image distillation LLM call failed for %s: %s", name, e,
        )

    # Vision call failed — emit a neutral stub. Important to still
    # return *something* readable so attachment_distilled / memory
    # blocks don't render an empty fence to the agent later.
    return (
        f"[Image — {name} ({mime}, {size_bytes} bytes). "
        "Vision caption distill unavailable; the image was still "
        "shown to the model on the current turn.]",
        "vision-caption+fallback",
    )
